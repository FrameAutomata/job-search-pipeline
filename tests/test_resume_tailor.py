"""Tests for pipeline/resume_tailor.py — slot classification, grounding,
budget fitting, prose rules, the verify/retry loop (renderer faked), caching
semantics, and threshold gating.

A synthetic docx is built in-test mirroring the real resume's structure:
name, contact, section headers, "Label: values" skills lines (bold label run),
'List Paragraph' bullets, bold "Company\\tDate" lines, and an education block."""

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402

from pipeline import resume_tailor as rt
from pipeline.apply import _should_tailor, _stem_matches_company
from pipeline.apply.queue import ApplyJob


def _build_resume(path: Path) -> Path:
    d = Document()
    d.add_paragraph("Jane Dev")                                      # 0 name
    d.add_paragraph("+1 555 | jane@x.com | Dallas, TX")              # 1 contact
    d.add_paragraph("Professional Summary")                          # 2 header
    d.add_paragraph("Engineer with 3 yrs across backend and DevOps " # 3 summary
                    "building production systems.")
    d.add_paragraph("Skills")                                        # 4 header
    p = d.add_paragraph()                                            # 5 skills
    r = p.add_run("Languages: "); r.bold = True
    p.add_run("Java, Python, Go, SQL")
    p2 = d.add_paragraph()                                           # 6 skills
    r2 = p2.add_run("Cloud / DevOps: "); r2.bold = True
    p2.add_run("AWS, Docker, Kubernetes, Terraform")
    d.add_paragraph("Professional Experience")                       # 7 header
    comp = d.add_paragraph()                                         # 8 company line
    rc = comp.add_run("Acme Corp\tJan 2023 – Present"); rc.bold = True
    role = d.add_paragraph()                                         # 9 role line
    rr = role.add_run("Software Engineer\tRemote"); rr.italic = True
    d.add_paragraph("Built REST APIs in Java handling 1M requests "  # 10 bullet
                    "daily with full CI/CD on AWS ECS.", style="List Paragraph")
    d.add_paragraph("Deployed Dockerized services to AWS ECS.",      # 11 bullet
                    style="List Paragraph")
    d.add_paragraph("Education & Certifications")                    # 12 header
    edu = d.add_paragraph()                                          # 13 education
    re_ = edu.add_run("State University\t2018 – 2022"); re_.bold = True
    d.add_paragraph("AWS Certified Cloud Practitioner, 2022",        # 14 cert bullet
                    style="List Paragraph")
    d.save(str(path))
    return path


@pytest.fixture
def resume_docx(tmp_path):
    return _build_resume(tmp_path / "resume.docx")


def _job(company="Acme", score=4.5, report="", role="Backend Engineer"):
    return ApplyJob(num="1", company=company, role=role,
                    url="u", score=score, report_path=report)


# ── classification ───────────────────────────────────────────────────────────

class TestExtractSlots:
    def test_classifies_only_editable_slots(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        kinds = {s.id: s.kind for s in slots}
        # summary (3), skills (5, 6), bullets (10, 11) — NOT the education
        # cert bullet (14), NOT name/contact/headers/company lines.
        assert kinds == {"s3": "summary", "s5": "skills", "s6": "skills",
                         "s10": "bullet", "s11": "bullet"}

    def test_education_bullets_never_editable(self, resume_docx):
        # A reworded credential is a fabrication — cert/degree bullets under
        # Education are protected even though they carry the bullet style.
        doc = Document(str(resume_docx))
        assert all(s.para_index != 14 for s in rt.extract_slots(doc))

    def test_unknown_header_ends_summary_section(self, tmp_path):
        # 'Core Competencies' isn't in _HEADERS — without the heading boundary,
        # it and every line under it leaked to the LLM as rewritable summary.
        d = Document()
        d.add_paragraph("Summary")
        d.add_paragraph("Engineer with 3 yrs building systems for production teams.")
        d.add_paragraph("Core Competencies")
        d.add_paragraph("Leadership and mentoring across multiple teams and offices.")
        p = tmp_path / "odd.docx"
        d.save(str(p))
        slots = rt.extract_slots(Document(str(p)))
        assert [s.para_index for s in slots] == [1]   # only the real summary

    def test_bullet_via_numbering_without_english_style(self, tmp_path):
        # Non-English Word writes localized style names ('Listenabsatz') — the
        # locale-independent signal is w:numPr.
        d = Document()
        d.add_paragraph("Experience")
        p = d.add_paragraph("Built data pipelines in Python for nightly loads.")
        numPr = OxmlElement("w:numPr")
        p._p.get_or_add_pPr().append(numPr)
        f = tmp_path / "locale.docx"
        d.save(str(f))
        slots = rt.extract_slots(Document(str(f)))
        assert [(s.kind, s.para_index) for s in slots] == [("bullet", 1)]

    def test_hyperlink_paragraph_protected(self, tmp_path):
        # p.text includes hyperlink text but p.runs doesn't — patching runs
        # would glue the link text onto the rewrite. Such paragraphs are
        # protected outright.
        d = Document()
        d.add_paragraph("Projects")
        p = d.add_paragraph("Built a CLI tool — ", style="List Paragraph")
        link = OxmlElement("w:hyperlink")
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "github.com/jane/tool"
        run.append(t)
        link.append(run)
        p._p.append(link)
        f = tmp_path / "link.docx"
        d.save(str(f))
        doc = Document(str(f))
        assert doc.paragraphs[1].text.endswith("github.com/jane/tool")  # sanity
        assert rt.extract_slots(doc) == []

    def test_skills_slot_excludes_bold_label(self, resume_docx):
        doc = Document(str(resume_docx))
        slot = next(s for s in rt.extract_slots(doc) if s.id == "s5")
        assert slot.label.strip() == "Languages:"
        assert slot.text == "Java, Python, Go, SQL"

    def test_nothing_editable_before_first_header(self, tmp_path):
        d = Document()
        d.add_paragraph("Stray bullet before any header", style="List Paragraph")
        p = tmp_path / "odd.docx"
        d.save(str(p))
        assert rt.extract_slots(Document(str(p))) == []


# ── patching ─────────────────────────────────────────────────────────────────

class TestApplyReplacements:
    def test_patches_within_budget_and_preserves_protected(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        changed, rejected = rt.apply_replacements(doc, slots, {
            "s10": "Built Java REST APIs at 1M req/day; full CI/CD.",
            "s5": "Go, Java, Python, SQL",
        })
        assert changed == 2 and rejected == []
        assert doc.paragraphs[10].text == "Built Java REST APIs at 1M req/day; full CI/CD."
        assert doc.paragraphs[5].runs[0].text == "Languages: "
        assert doc.paragraphs[5].runs[0].bold is True
        assert doc.paragraphs[5].text == "Languages: Go, Java, Python, SQL"
        assert doc.paragraphs[0].text == "Jane Dev"
        assert doc.paragraphs[8].text == "Acme Corp\tJan 2023 – Present"

    def test_skills_separator_space_preserved(self, tmp_path):
        # When the label run has NO trailing space (the space lives in the
        # values run), the patch must keep a separator — 'Languages:Go, Java'
        # jammed every skills line against its label.
        d = Document()
        d.add_paragraph("Skills")
        p = d.add_paragraph()
        r = p.add_run("Languages:"); r.bold = True
        p.add_run(" Java, Python, Go")
        f = tmp_path / "space.docx"
        d.save(str(f))
        doc = Document(str(f))
        slots = rt.extract_slots(doc)
        changed, _ = rt.apply_replacements(doc, slots, {slots[0].id: "Go, Java, Python"})
        assert changed == 1
        assert doc.paragraphs[1].text == "Languages: Go, Java, Python"

    def test_prose_patch_uses_dominant_run_formatting(self, tmp_path):
        # A bullet with a short bold lead-in must not become ALL bold after the
        # rewrite — the replacement goes into the dominant (longest) run.
        d = Document()
        d.add_paragraph("Experience")
        p = d.add_paragraph(style="List Paragraph")
        lead = p.add_run("Led "); lead.bold = True
        p.add_run("the migration of twelve services to Kubernetes with zero downtime.")
        f = tmp_path / "bold.docx"
        d.save(str(f))
        doc = Document(str(f))
        slots = rt.extract_slots(doc)
        changed, _ = rt.apply_replacements(doc, slots, {slots[0].id:
                                           "Migrated twelve services to Kubernetes with zero downtime."})
        assert changed == 1
        para = doc.paragraphs[1]
        assert para.text == "Migrated twelve services to Kubernetes with zero downtime."
        target = next(r for r in para.runs if r.text)
        assert target.bold is not True   # carried the dominant (non-bold) format

    def test_unfittable_replacement_rejected(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        original = doc.paragraphs[11].text
        too_long = "x" * (next(s.max_chars for s in slots if s.id == "s11") + 50)
        changed, rejected = rt.apply_replacements(doc, slots, {"s11": too_long})
        assert changed == 0 and rejected == ["s11"]
        assert doc.paragraphs[11].text == original

    def test_over_budget_prose_trimmed_at_clause_boundary(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        s3 = next(s for s in slots if s.id == "s3")
        first = "Backend-focused engineer with 3 yrs building production systems."
        new = first + " " + "Extra trailing sentence that pushes well past the budget limit." * 3
        assert len(new) > s3.max_chars
        changed, rejected = rt.apply_replacements(doc, slots, {"s3": new})
        assert changed == 1 and rejected == []
        assert doc.paragraphs[3].text.startswith("Backend-focused engineer")
        assert len(doc.paragraphs[3].text) <= s3.max_chars

    def test_skills_fabricated_tokens_dropped(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        resume_text = "\n".join(p.text for p in doc.paragraphs)
        changed, rejected = rt.apply_replacements(
            doc, slots, {"s5": "Python, JavaScript, Go, Java, SQL"}, resume_text)
        assert changed == 1
        assert doc.paragraphs[5].text == "Languages: Python, Go, Java, SQL"

    def test_skills_all_fabricated_rejected(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        resume_text = "\n".join(p.text for p in doc.paragraphs)
        original = doc.paragraphs[5].text
        changed, rejected = rt.apply_replacements(
            doc, slots, {"s5": "Rust, Elixir, Haskell"}, resume_text)
        assert changed == 0 and rejected == ["s5"]
        assert doc.paragraphs[5].text == original

    def test_unknown_or_unchanged_slots_ignored(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        same = doc.paragraphs[10].text
        changed, rejected = rt.apply_replacements(doc, slots, {"s10": same, "s999": "x", "s10x": 7})
        assert changed == 0 and rejected == []

    def test_bullet_style_survives_patch(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        rt.apply_replacements(doc, slots, {"s10": "Shorter bullet."})
        assert doc.paragraphs[10].style.name == "List Paragraph"

    def test_summary_gets_extra_headroom(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        s3 = next(s for s in slots if s.kind == "summary")
        s10 = next(s for s in slots if s.id == "s10")
        assert s3.max_chars > int(len(s3.text) * 1.25)
        assert s10.max_chars <= int(len(s10.text) * 1.25) + 5


# ── grounding semantics ──────────────────────────────────────────────────────

class TestSkillGrounding:
    RESUME = "Java and Python with CI/CD pipelines, Linux admin, AWS ECS deploys."

    def test_verbatim_compound_passes(self):
        assert rt._skill_in_resume("CI/CD", self.RESUME) is True

    def test_slash_smuggling_blocked(self):
        # 'CD' alone matches inside 'CI/CD' ('/' is a non-word char) — any-part
        # grounding let fabricated halves ride along. All parts must ground.
        assert rt._skill_in_resume("Jenkins/CD", self.RESUME) is False
        assert rt._skill_in_resume("Azure/AWS", self.RESUME) is False
        assert rt._skill_in_resume("Java/Kotlin", self.RESUME) is False

    def test_slash_all_parts_grounded_passes(self):
        assert rt._skill_in_resume("Java/Python", self.RESUME) is True

    def test_paren_semantics(self):
        assert rt._skill_in_resume("AWS (ECS)", self.RESUME) is True
        assert rt._skill_in_resume("AWS (ECS, DynamoDB)", self.RESUME) is False


# ── prompt / prose rules / parsing ───────────────────────────────────────────

class TestBuildPrompt:
    def test_includes_rules_slots_and_context(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        system, user = rt.build_prompt(slots, "FULL RESUME TEXT", _job(report=""),
                                       "REPORT NOTES")
        assert "Never invent" in system and "max_chars" in system
        assert "ALWAYS rewrite" in system and "FAILURE" in system
        assert "PROSE HONESTY" in system and "NEVER with the target" in system
        # JD/report framed as data, not instructions (prompt-injection guard).
        assert "never instructions" in system
        assert "Acme — Backend Engineer" in user
        assert "REPORT NOTES" in user and "FULL RESUME TEXT" in user
        assert '"s10"' in user

    def test_jd_text_included_when_present(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        _, user = rt.build_prompt(slots, "cv", _job(), "", jd_text="THE ACTUAL JD")
        assert "JOB DESCRIPTION" in user and "THE ACTUAL JD" in user
        assert "untrusted" in user

    def test_feedback_appended_to_system(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        system, _ = rt.build_prompt(slots, "cv", _job(), "",
                                    feedback=rt._OVERFLOW_FEEDBACK)
        assert "OVERFLOWED" in system
        system2, _ = rt.build_prompt(slots, "cv", _job(), "",
                                     feedback=rt._TITLE_FEEDBACK)
        assert "REJECTED" in system2


class TestProseRules:
    def test_retitled_summary_dropped(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        job = _job(role="Application Support Engineer")
        reps = {"s3": "Application Support Engineer with production experience.",
                "s10": "Fine bullet."}
        out, notes = rt.enforce_prose_rules(reps, slots, job)
        assert "s3" not in out and "s10" in out
        assert notes and "target job title" in notes[0]

    def test_embedded_words_are_not_a_retitle(self, resume_docx):
        # Token-sequence matching: 'big data engineering' must NOT match a
        # 'Data Engineer' role (substring matching flagged it and burned the
        # corrective retry on a false positive).
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        job = _job(role="Data Engineer")
        reps = {"s3": "Engineer with 3 yrs building big data engineering pipelines."}
        out, notes = rt.enforce_prose_rules(reps, slots, job)
        assert "s3" in out and notes == []

    def test_tenure_misattribution_dropped(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        reps = {"s3": "Engineer with 3+ years at Acme Corp building systems."}
        out, notes = rt.enforce_prose_rules(reps, slots, _job())
        assert "s3" not in out
        assert notes and "duration to an employer" in notes[0]

    def test_seniority_label_dropped(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        reps = {"s10": "Senior engineer driving Java REST APIs at 1M req/day."}
        out, notes = rt.enforce_prose_rules(reps, slots, _job())
        assert "s10" not in out
        assert notes and "seniority label" in notes[0]

    def test_matching_original_opener_allowed(self, resume_docx):
        doc = Document(str(resume_docx))
        slots = rt.extract_slots(doc)
        job = _job(role="Backend Engineer")
        reps = {"s3": "Engineer with 3 yrs building backend and DevOps systems for production."}
        out, notes = rt.enforce_prose_rules(reps, slots, job)
        assert "s3" in out and notes == []

    def test_leads_with_role_title_helper(self):
        assert rt._leads_with_role_title(
            "Application Support Engineer with experience...",
            "Software engineer (~3 yrs) across backend...",
            "Application Support Engineer") is True
        assert rt._leads_with_role_title(
            "Software engineer (~3 yrs) focused on support tooling...",
            "Software engineer (~3 yrs) across backend...",
            "Software Engineer") is False
        # word-embedding false positives (substring matching regression)
        assert rt._leads_with_role_title(
            "Engineer shipping the latest engineering tooling daily...",
            "Software engineer (~3 yrs)...", "Test Engineer") is False


class TestParseReplacements:
    def test_plain_json(self):
        assert rt.parse_replacements('{"s3": "new"}') == {"s3": "new"}

    def test_fenced_with_prose(self):
        raw = 'Here you go:\n```json\n{"s3": "new", "s5": "vals"}\n```\nDone.'
        assert rt.parse_replacements(raw) == {"s3": "new", "s5": "vals"}

    def test_trailing_prose_with_braces(self):
        # The greedy first-{-to-last-} slice broke whenever trailing prose
        # contained a '}' — the shared loose parser handles it.
        raw = '{"s3": "new"}\nNote: I kept the {descriptor} as required.}'
        assert rt.parse_replacements(raw) == {"s3": "new"}

    def test_garbage_and_non_string_values(self):
        assert rt.parse_replacements("no json here") == {}
        assert rt.parse_replacements('{"s3": 42, "s5": "ok"}') == {"s5": "ok"}


# ── naming / matching ────────────────────────────────────────────────────────

class TestResumePaths:
    def test_dotted_company_names_intact(self, tmp_path):
        # .with_suffix() truncated at the company's last dot ('St. Jude Medical
        # - resume' → 'St.pdf') and collided across companies.
        docx_p, pdf_p = rt.resume_paths(tmp_path, "St. Jude Medical")
        assert docx_p.name == "St. Jude Medical - resume.docx"
        assert pdf_p.name == "St. Jude Medical - resume.pdf"
        other_docx, _ = rt.resume_paths(tmp_path, "St. David's Healthcare")
        assert other_docx.name != docx_p.name


class TestStemMatchesCompany:
    def test_contiguous_token_match(self):
        assert _stem_matches_company("CV - Apexon Inc", "apexon") is True
        assert _stem_matches_company("Apexon - resume", "apexon") is True
        assert _stem_matches_company("CV - Apexon Inc", "apexoninc") is True

    def test_substring_of_token_does_not_match(self):
        # 'Meta' must not match 'Metabase - resume.pdf' (another company).
        assert _stem_matches_company("Metabase - resume", "meta") is False
        assert _stem_matches_company("Asana - resume", "sana") is False


# ── generate loop ────────────────────────────────────────────────────────────

class TestGenerateForJob:
    """End-to-end with a fake caller and faked renderer (no LibreOffice)."""

    def _setup(self, tmp_path, resume_docx, monkeypatch, *, pages_seq, baseline=1,
               soffice=True):
        co = tmp_path / "career-ops"
        (co / "output").mkdir(parents=True)
        monkeypatch.setattr(rt, "source_docx", lambda: resume_docx)
        monkeypatch.setattr(rt, "_baseline_pages", lambda src: baseline)
        monkeypatch.setattr(rt, "find_soffice",
                            (lambda: "soffice") if soffice else (lambda: None))
        seq = iter(pages_seq)

        def fake_render(docx_path, out_dir):
            pdf = Path(out_dir) / (Path(docx_path).stem + ".pdf")
            pdf.write_bytes(b"%PDF-fake")
            return pdf
        monkeypatch.setattr(rt, "render_pdf", fake_render)
        monkeypatch.setattr(rt, "page_count", lambda p: next(seq))
        return co

    def test_success_returns_verified_pdf(self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        out = rt.generate_for_job(
            co, _job(), caller=lambda s, u: '{"s10": "Java APIs, tailored."}')
        assert out is not None and out.name == "Acme - resume.pdf" and out.exists()
        saved = Document(str(co / "output" / "Acme - resume.docx"))
        assert saved.paragraphs[10].text == "Java APIs, tailored."
        # no work files left behind
        assert not list((co / "output").glob("*.work.*"))

    def test_overflow_retries_then_falls_back(self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[2, 2])
        calls = []
        def caller(system, user):
            calls.append(system)
            return '{"s10": "Tailored but somehow too long."}'
        assert rt.generate_for_job(co, _job(), caller=caller) is None
        assert len(calls) == 2 and "OVERFLOWED" in calls[1]
        # nothing cached, nothing left over
        assert not list((co / "output").glob("*"))

    def test_overflow_then_fits_on_retry(self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[2, 1])
        out = rt.generate_for_job(
            co, _job(), caller=lambda s, u: '{"s10": "Short."}')
        assert out is not None and out.suffix == ".pdf"

    def test_render_failure_with_soffice_present_returns_none(
            self, tmp_path, resume_docx, monkeypatch):
        # Renderer installed but conversion failed: never ship/cache unverified.
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[])
        monkeypatch.setattr(rt, "render_pdf", lambda d, o: None)
        out = rt.generate_for_job(
            co, _job(), caller=lambda s, u: '{"s10": "Tailored."}')
        assert out is None
        assert not list((co / "output").glob("*"))   # no unverified docx cached

    def test_no_renderer_returns_docx(self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[],
                         soffice=False, baseline=None)
        out = rt.generate_for_job(
            co, _job(), caller=lambda s, u: '{"s10": "Tailored."}')
        assert out is not None and out.suffix == ".docx"

    def test_existing_reused_without_llm(self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[])
        (co / "output" / "Acme - resume.pdf").write_bytes(b"%PDF-existing")
        def boom(s, u):
            raise AssertionError("LLM must not be called for a cached resume")
        out = rt.generate_for_job(co, _job(), caller=boom)
        assert out is not None and out.name == "Acme - resume.pdf"

    def test_role_mismatch_regenerates(self, tmp_path, resume_docx, monkeypatch):
        # Cache tailored for a different role at the same company is stale.
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        import shutil as _sh
        _sh.copyfile(resume_docx, co / "output" / "Acme - resume.docx")
        d = Document(str(co / "output" / "Acme - resume.docx"))
        d.core_properties.subject = "DevOps Engineer"
        d.save(str(co / "output" / "Acme - resume.docx"))
        (co / "output" / "Acme - resume.pdf").write_bytes(b"%PDF-old-role")
        calls = []
        def caller(s, u):
            calls.append(1)
            return '{"s10": "Backend-tailored."}'
        out = rt.generate_for_job(co, _job(role="Backend Engineer"), caller=caller)
        assert calls and out is not None      # regenerated, not reused

    def test_hand_edited_docx_wins_over_stale_pdf(self, tmp_path, resume_docx, monkeypatch):
        import os as _os
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        docx_p = co / "output" / "Acme - resume.docx"
        pdf_p = co / "output" / "Acme - resume.pdf"
        import shutil as _sh
        _sh.copyfile(resume_docx, docx_p)
        pdf_p.write_bytes(b"%PDF-stale")
        now = docx_p.stat().st_mtime
        _os.utime(pdf_p, (now - 100, now - 100))   # pdf older than docx
        def boom(s, u):
            raise AssertionError("hand-edit path must not call the LLM")
        out = rt.generate_for_job(co, _job(role=""), caller=boom)
        assert out == pdf_p
        assert pdf_p.read_bytes() == b"%PDF-fake"   # re-rendered from the docx

    def test_docx_only_cache_healed_when_renderer_appears(
            self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        import shutil as _sh
        _sh.copyfile(resume_docx, co / "output" / "Acme - resume.docx")
        def boom(s, u):
            raise AssertionError("healing must not call the LLM")
        out = rt.generate_for_job(co, _job(role=""), caller=boom)
        assert out is not None and out.suffix == ".pdf"   # rendered + verified

    def test_no_source_docx_returns_none(self, tmp_path, monkeypatch):
        co = tmp_path / "career-ops"
        (co / "output").mkdir(parents=True)
        monkeypatch.setattr(rt, "source_docx", lambda: None)
        assert rt.generate_for_job(co, _job(), caller=lambda s, u: "{}") is None

    def test_no_change_returns_none(self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        assert rt.generate_for_job(co, _job(), caller=lambda s, u: "{}") is None

    def test_retitled_summary_retried_with_feedback_then_applied(
            self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        job = _job(role="Application Support Engineer")
        systems = []
        def caller(system, user):
            systems.append(system)
            if len(systems) == 1:
                return '{"s3": "Application Support Engineer with production experience."}'
            return '{"s3": "Engineer with 3 yrs of production support across financial systems."}'
        out = rt.generate_for_job(co, job, caller=caller)
        assert out is not None
        assert len(systems) == 2 and "REJECTED" in systems[1]
        saved = Document(str(co / "output" / "Acme - resume.docx"))
        assert saved.paragraphs[3].text.startswith("Engineer with 3 yrs")

    def test_title_retry_carries_validated_rewrites(
            self, tmp_path, resume_docx, monkeypatch):
        # Attempt-0's good bullet must survive even if attempt-1 returns only
        # the corrected summary (carry-merge: no re-roll of accepted slots).
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        job = _job(role="Application Support Engineer")
        n = []
        def caller(system, user):
            n.append(1)
            if len(n) == 1:
                return ('{"s3": "Application Support Engineer with experience.", '
                        '"s10": "Validated bullet from attempt zero."}')
            return '{"s3": "Engineer with 3 yrs of support work."}'
        out = rt.generate_for_job(co, job, caller=caller)
        assert out is not None
        saved = Document(str(co / "output" / "Acme - resume.docx"))
        assert saved.paragraphs[10].text == "Validated bullet from attempt zero."
        assert saved.paragraphs[3].text.startswith("Engineer with 3 yrs")

    def test_retitled_twice_drops_summary_keeps_rest(
            self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        job = _job(role="Application Support Engineer")
        original_summary = Document(str(resume_docx)).paragraphs[3].text
        def caller(system, user):
            return ('{"s3": "Application Support Engineer with production experience.", '
                    '"s10": "Built Java REST APIs; production support focus."}')
        out = rt.generate_for_job(co, job, caller=caller)
        assert out is not None
        saved = Document(str(co / "output" / "Acme - resume.docx"))
        assert saved.paragraphs[3].text == original_summary
        assert saved.paragraphs[10].text == "Built Java REST APIs; production support focus."

    def test_metadata_recorded(self, tmp_path, resume_docx, monkeypatch):
        co = self._setup(tmp_path, resume_docx, monkeypatch, pages_seq=[1])
        rt.generate_for_job(co, _job(role="Backend Engineer"),
                            caller=lambda s, u: '{"s10": "Tailored."}')
        saved = Document(str(co / "output" / "Acme - resume.docx"))
        assert saved.core_properties.title == ""
        assert saved.core_properties.revision == 1
        assert saved.core_properties.subject == "Backend Engineer"


class TestShouldTailor:
    def test_gates_on_score(self):
        assert _should_tailor(_job(score=4.0), 4.0) is True
        assert _should_tailor(_job(score=4.6), 4.0) is True
        assert _should_tailor(_job(score=3.9), 4.0) is False
        assert _should_tailor(_job(score=None), 4.0) is False   # --apply-url one-off


class TestSourceDocx:
    def test_env_relative_anchored_to_root(self, tmp_path, monkeypatch):
        target = rt.ROOT / "resumes" / "resume.docx"   # exists in this repo
        monkeypatch.setenv("RESUME_DOCX_PATH", "resumes/resume.docx")
        monkeypatch.chdir(tmp_path)                    # cwd must not matter
        if target.is_file():
            assert rt.source_docx() == rt.ROOT / "resumes" / "resume.docx"

    def test_env_missing_warns_and_returns_none(self, tmp_path, monkeypatch, capsys):
        monkeypatch.setenv("RESUME_DOCX_PATH", str(tmp_path / "nope.docx"))
        assert rt.source_docx() is None
        assert "RESUME_DOCX_PATH" in capsys.readouterr().out


class TestSofficeDiscovery:
    def test_env_override_wins(self, tmp_path, monkeypatch):
        fake = tmp_path / "soffice.exe"
        fake.write_bytes(b"")
        monkeypatch.setenv("SOFFICE_PATH", str(fake))
        assert rt.find_soffice() == str(fake)


class TestJdText:
    def test_prefers_local_jds_file(self, tmp_path):
        from pipeline.cover_letters import jd_text_for_job
        co = tmp_path / "career-ops"
        (co / "batch" / "jds").mkdir(parents=True)
        (co / "batch" / "jds" / "7.txt").write_text("CACHED JD", encoding="utf-8")
        job = ApplyJob(num="7", company="X", role="Y", url="u", score=4.0)
        assert jd_text_for_job(co, None, job) == "CACHED JD"

    def test_falls_back_to_artifact_then_fetch(self, tmp_path, monkeypatch):
        from pipeline.cover_letters import jd_text_for_job
        co = tmp_path / "career-ops"
        co.mkdir()
        art = tmp_path / "artifact"
        (art / "batch" / "jds").mkdir(parents=True)
        (art / "batch" / "jds" / "7.txt").write_text("ARTIFACT JD", encoding="utf-8")
        job = ApplyJob(num="7", company="X", role="Y", url="u", score=4.0)
        assert jd_text_for_job(co, art, job) == "ARTIFACT JD"
        job2 = ApplyJob(num="8", company="X", role="Y",
                        url="https://www.linkedin.com/jobs/view/123/", score=4.0)
        monkeypatch.setattr("pipeline.screen.fetch_and_classify",
                            lambda url, timeout=8: ("active", "", "<body>FETCHED JD</body>"))
        assert "FETCHED JD" in jd_text_for_job(co, art, job2)

    def test_empty_when_nothing_available(self, tmp_path):
        from pipeline.cover_letters import jd_text_for_job
        co = tmp_path / "career-ops"
        co.mkdir()
        job = ApplyJob(num="", company="X", role="Y", url="https://example.com/x", score=4.0)
        assert jd_text_for_job(co, None, job) == ""
