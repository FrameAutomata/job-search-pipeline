"""Tests for pipeline.resume_text — format-dispatching resume extraction.

The module is the single source of truth for turning a resume file (PDF / DOCX /
ODT / TXT) into plain text, used by both the keyword-scoring filter and the UI
onboarding upload. These tests pin the DOCX/ODT/TXT branches, the dispatch
contract, and the PDF run-together retry.
"""

import sys
import types
from pathlib import Path

import pytest

from pipeline import resume_text


def _make_docx(path: Path, paragraphs, table=None) -> Path:
    from docx import Document
    d = Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table:
        t = d.add_table(rows=1, cols=len(table))
        for i, cell in enumerate(table):
            t.rows[0].cells[i].text = cell
    d.save(str(path))
    return path


def _make_odt(path: Path, headings_and_paras) -> Path:
    """headings_and_paras: list of ("h", text) | ("p", text) in document order."""
    from odf.opendocument import OpenDocumentText
    from odf.text import H, P
    od = OpenDocumentText()
    for kind, text in headings_and_paras:
        if kind == "h":
            od.text.addElement(H(outlinelevel=1, text=text))
        else:
            od.text.addElement(P(text=text))
    od.save(str(path))
    return path


class TestExtractResumeText:
    def test_docx_paragraphs_and_tables(self, tmp_path):
        """DOCX extraction returns paragraph text AND table-cell text (skills are
        frequently laid out in a table)."""
        f = _make_docx(
            tmp_path / "resume.docx",
            ["Jane Dev", "jane@example.com | Dallas, TX", "Python, AWS, Go"],
            table=["React", "Kubernetes"],
        )
        text = resume_text.extract_resume_text(f)
        assert "Jane Dev" in text
        assert "Python, AWS, Go" in text
        assert "React" in text and "Kubernetes" in text

    def test_odt_preserves_heading_and_paragraph_order(self, tmp_path):
        """ODT extraction walks the body in document order so the name (a heading)
        leads — the contact-info parser keys off the first line."""
        f = _make_odt(
            tmp_path / "resume.odt",
            [("h", "Jane Dev"), ("p", "jane@example.com | Dallas, TX"),
             ("p", "Python, AWS, Go")],
        )
        text = resume_text.extract_resume_text(f)
        lines = [ln for ln in text.splitlines() if ln.strip()]
        assert lines[0] == "Jane Dev"
        assert "Python, AWS, Go" in text

    def test_txt_passthrough(self, tmp_path):
        f = tmp_path / "resume.txt"
        f.write_text("Jane Dev\nPython", encoding="utf-8")
        assert resume_text.extract_resume_text(f) == "Jane Dev\nPython"

    def test_suffix_is_case_insensitive(self, tmp_path):
        f = _make_docx(tmp_path / "resume.DOCX", ["Hello world"])
        assert "Hello world" in resume_text.extract_resume_text(f)

    def test_unsupported_suffix_raises_valueerror(self, tmp_path):
        f = tmp_path / "resume.rtf"
        f.write_text("x", encoding="utf-8")
        with pytest.raises(ValueError):
            resume_text.extract_resume_text(f)

    def test_pdf_routes_to_pdf_extractor(self, tmp_path, monkeypatch):
        """A .pdf suffix dispatches to the pdfplumber branch (patched here so the
        test needs no real PDF). Pins the dispatch, not pdfplumber itself."""
        called = {}

        def fake_pdf(p):
            called["p"] = p
            return "PDF TEXT"

        monkeypatch.setattr(resume_text, "_from_pdf", fake_pdf)
        f = tmp_path / "resume.pdf"
        f.write_bytes(b"%PDF-1.4 fake")
        assert resume_text.extract_resume_text(f) == "PDF TEXT"
        assert called["p"] == f


class TestExtractResumeBytes:
    def test_dispatches_on_filename_suffix(self, tmp_path):
        """extract_resume_bytes routes by the *filename's* suffix (the upload has
        bytes + a client filename, not a path on disk)."""
        f = _make_docx(tmp_path / "src.docx", ["Jane Dev", "Python, AWS"])
        data = f.read_bytes()
        text = resume_text.extract_resume_bytes(data, "whatever-the-user-named-it.docx")
        assert "Jane Dev" in text and "Python, AWS" in text

    def test_unsupported_filename_raises(self):
        with pytest.raises(ValueError):
            resume_text.extract_resume_bytes(b"x", "resume.rtf")


class TestSupportedSuffixes:
    def test_advertises_the_three_import_formats_plus_txt(self):
        s = set(resume_text.SUPPORTED_SUFFIXES)
        assert {".pdf", ".docx", ".odt", ".txt"} <= s


# Enough tokens to clear _looks_run_together's 20-token floor.
_CLEAN = " ".join(
    ["Managed", "high-volume", "inbound", "patient", "calls", "and", "explained",
     "insurance", "benefits", "while", "coordinating", "appointment",
     "scheduling", "with", "provider", "offices", "to", "verify", "network",
     "eligibility", "for", "seamless", "care", "coordination", "standards"]
)
# The same prose as one glyph run per line — what pdfplumber emits when the
# PDF's kerning is tighter than the default x_tolerance.
_RUNON = "\n".join(
    ["Managedhigh-volumeinboundpatientcallsandexplainedinsurancebenefits",
     "whilecoordinatingappointmentschedulingwithprovideroffices",
     "toverifynetworkeligibilityforseamlesscarecoordinationstandards"] * 9
)


class _FakePage:
    """Returns different text per x_tolerance, recording what it was asked for."""

    def __init__(self, by_tolerance, calls):
        self._by_tolerance = by_tolerance
        self._calls = calls

    def extract_text(self, **kwargs):
        tol = kwargs.get("x_tolerance")
        self._calls.append(tol)
        return self._by_tolerance[tol]


class _FakePDF:
    def __init__(self, pages):
        self.pages = pages

    def __enter__(self):
        return self

    def __exit__(self, *exc):
        return False


def _fake_pdfplumber(monkeypatch, by_tolerance):
    """Install a stub pdfplumber and return the list of x_tolerance values it
    was called with. _from_pdf imports pdfplumber lazily inside the function,
    so patching sys.modules is what reaches it."""
    calls = []
    page = _FakePage(by_tolerance, calls)
    monkeypatch.setitem(
        sys.modules, "pdfplumber",
        types.SimpleNamespace(open=lambda path: _FakePDF([page])),
    )
    return calls


class TestLooksRunTogether:
    """The heuristic gating the retry."""

    def test_clean_prose_is_not_run_together(self):
        assert resume_text._looks_run_together(_CLEAN) is False

    def test_glyph_run_is_run_together(self):
        assert resume_text._looks_run_together(_RUNON) is True

    def test_short_text_is_never_judged(self):
        """Under the token floor we decline to guess — a near-empty extraction
        is a different problem, and guessing would add a second failure mode."""
        assert resume_text._looks_run_together("Averyveryverylongsingletoken") is False


class TestFromPdfRetry:
    def test_clean_pdf_is_not_retried(self, monkeypatch):
        """A PDF that extracts fine must not be re-extracted — a tighter
        tolerance can split words that were already correct."""
        calls = _fake_pdfplumber(monkeypatch, {None: _CLEAN})
        assert resume_text._from_pdf(Path("x.pdf")) == _CLEAN
        assert calls == [None], "retry attempted on a clean extraction"

    def test_retry_used_when_it_resolves_the_run(self, monkeypatch, capsys):
        """The bug this exists for: at default spacing the text has no word
        boundaries, so the tighter re-extraction is what callers get."""
        calls = _fake_pdfplumber(
            monkeypatch,
            {None: _RUNON, resume_text._PDF_RETRY_X_TOLERANCE: _CLEAN},
        )
        assert resume_text._from_pdf(Path("resume.pdf")) == _CLEAN
        assert calls == [None, resume_text._PDF_RETRY_X_TOLERANCE]
        assert "ran together" in capsys.readouterr().out

    def test_retry_discarded_when_it_does_not_help(self, monkeypatch, capsys):
        """A tighter tolerance that still yields run-together text is no
        evidence it helped, so keep pdfplumber's default output."""
        _fake_pdfplumber(
            monkeypatch,
            {None: _RUNON, resume_text._PDF_RETRY_X_TOLERANCE: _RUNON},
        )
        assert resume_text._from_pdf(Path("resume.pdf")) == _RUNON
        assert "ran together" not in capsys.readouterr().out


class TestResolveResumePath:
    """Which file counts as "the resume". Lives here rather than with the filter
    because the UI's setup wizard asks the same question through the same
    function and cannot import pipeline.filter (yake, pandas) — answering it
    twice is what let a RESUME_PATH resume read as absent in the wizard while
    the filter was scoring against it (#145)."""

    def _resumes_dir(self, tmp_path):
        d = tmp_path / "resumes"
        d.mkdir()
        return d

    def test_prefers_explicit_existing_path(self, tmp_path):
        d = self._resumes_dir(tmp_path)
        (d / "custom.docx").write_bytes(b"x")
        chosen = resume_text.resolve_resume_path(str(d / "custom.docx"), tmp_path)
        assert chosen == (d / "custom.docx")

    def test_explicit_path_is_honored_even_when_missing(self, tmp_path):
        """So the caller's not-found error names what the user pointed at,
        rather than silently falling back to a different resume."""
        chosen = resume_text.resolve_resume_path("docs/nope.docx", tmp_path)
        assert chosen == (tmp_path / "docs" / "nope.docx")

    def test_discovers_dropped_docx_when_unset(self, tmp_path):
        d = self._resumes_dir(tmp_path)
        (d / "resume.docx").write_bytes(b"x")
        chosen = resume_text.resolve_resume_path("", tmp_path)
        assert chosen == (d / "resume.docx")

    def test_discovers_dropped_odt_when_unset(self, tmp_path):
        d = self._resumes_dir(tmp_path)
        (d / "resume.odt").write_bytes(b"x")
        chosen = resume_text.resolve_resume_path("", tmp_path)
        assert chosen == (d / "resume.odt")

    def test_pdf_wins_over_docx_when_both_present(self, tmp_path):
        """Deterministic precedence — pdf first, matching the historical default."""
        d = self._resumes_dir(tmp_path)
        (d / "resume.pdf").write_bytes(b"x")
        (d / "resume.docx").write_bytes(b"x")
        chosen = resume_text.resolve_resume_path("", tmp_path)
        assert chosen == (d / "resume.pdf")

    def test_probe_names_track_the_import_formats(self):
        """Derived, not restated — adding a format can't leave the probe behind."""
        assert resume_text.PROBE_NAMES == tuple(
            f"resume{s}" for s in resume_text.IMPORT_SUFFIXES)
