"""Tests for pipeline/app/onboard.py.

Pure mapping/parsing logic runs everywhere; the node round-trip and pdfplumber
extraction are guarded so the suite still passes without node / pdfplumber.
"""

import base64
import re
import shutil
import subprocess
from pathlib import Path

import pytest
import yaml

from pipeline.app import onboard
from pipeline.sites import SUPPORTED_SITES


@pytest.fixture
def html():
    """The onboarding wizard's static markup, for the drift checks below."""
    root = Path(__file__).resolve().parent.parent
    return (root / "pipeline" / "app" / "static" / "onboard.html").read_text(encoding="utf-8")


def _node_deps_available() -> bool:
    """True only if node AND the npm packages setup-profile.mjs imports
    (yaml, pdf-parse) resolve. CI has node but doesn't `npm install`, so the
    round-trip test skips there; it runs locally where setup has installed deps."""
    if shutil.which("node") is None:
        return False
    repo = Path(__file__).resolve().parent.parent
    try:
        r = subprocess.run(
            ["node", "-e", "require.resolve('yaml'); require.resolve('pdf-parse')"],
            cwd=str(repo), capture_output=True, timeout=15,
        )
        return r.returncode == 0
    except Exception:
        return False


class TestParseLocations:
    @pytest.mark.parametrize("text,expected", [
        ("", []),
        ("Dallas, TX", ["Dallas, TX"]),
        ("US Remote, Dallas, TX, Fort Worth, TX",
         ["US Remote", "Dallas, TX", "Fort Worth, TX"]),
        ("Toronto, ON, Montreal, QC", ["Toronto, ON", "Montreal, QC"]),
        ("London, UK", ["London, UK"]),
        ("Remote", ["Remote"]),
    ])
    def test_keeps_city_state_pairs(self, text, expected):
        assert onboard.parse_locations(text) == expected


class TestInferCountry:
    @pytest.mark.parametrize("loc,country", [
        ("Dallas, TX", "USA"),
        ("Toronto, ON", "Canada"),
        ("London, UK", "UK"),
        ("Sydney, Australia", "Australia"),
        ("US Remote", "USA"),
        ("somewhere weird", "USA"),  # safe default
    ])
    def test_infer(self, loc, country):
        assert onboard.infer_country(loc) == country

    def test_remote_location_map(self):
        assert onboard.infer_remote_location("Canada") == "Canada"
        assert onboard.infer_remote_location("USA") == "United States"
        assert onboard.infer_remote_location("UK") == "United Kingdom"


class TestBuildOnboardingJson:
    def test_maps_form_and_marks_remote_vs_local(self):
        form = {
            "name": "Jane Dev",
            "target_roles": "Senior Full-Stack Engineer, Backend Engineer",
            "negative_roles": "Intern, Manager",
            "comp_target": "$150K",
            "locations": "US Remote, Dallas, TX",
            "distance": "40",
            "hours_old": "48",
            "results_wanted": "100",
            "sites": ["indeed", "linkedin"],
            "include_easy_apply": True,
            "eeo_gender": "Female",
            "eeo_disability": "No, I do not have a disability",
        }
        payload = onboard.build_onboarding_json(form, resume_text="resume body")
        assert payload["resumeText"] == "resume body"
        assert payload["info"]["name"] == "Jane Dev"
        assert payload["info"]["eeoGender"] == "Female"
        assert payload["info"]["eeoDisability"] == "No, I do not have a disability"
        assert payload["info"]["eeoVeteran"] == ""        # unset → blank → declines
        assert payload["criteria"]["targetRoles"] == ["Senior Full-Stack Engineer", "Backend Engineer"]
        assert payload["criteria"]["negativeRoles"] == ["Intern", "Manager"]
        locs = payload["searchSettings"]["locations"]
        assert locs[0]["isRemote"] is True and locs[0]["location"] == "United States"
        assert locs[1]["isRemote"] is False and locs[1]["distance"] == 40
        assert payload["searchSettings"]["hoursOld"] == 48
        assert payload["searchSettings"]["includeEasyApply"] is True

    def test_defaults_when_empty(self):
        payload = onboard.build_onboarding_json({}, resume_text="")
        # No locations -> a single US-Remote default entry.
        locs = payload["searchSettings"]["locations"]
        assert len(locs) == 1 and locs[0]["isRemote"] is True
        assert payload["searchSettings"]["sites"] == ["indeed", "linkedin"]
        assert payload["searchSettings"]["resultsWanted"] == 100

    def test_unsupported_sites_are_filtered_out(self):
        # A stale saved wizard state (or hand-crafted POST) may still carry the
        # retired boards — they must not reach the generated search config.
        payload = onboard.build_onboarding_json(
            {"sites": ["indeed", "glassdoor", "zip_recruiter", "google"]}, resume_text=""
        )
        assert payload["searchSettings"]["sites"] == ["indeed"]

    def test_only_unsupported_sites_falls_back_to_default(self):
        payload = onboard.build_onboarding_json(
            {"sites": ["glassdoor", "google"]}, resume_text=""
        )
        assert payload["searchSettings"]["sites"] == ["indeed", "linkedin"]

    def test_maps_indeed_consent_toggles(self):
        info = onboard.build_onboarding_json(
            {"data_processing_consent": "no", "save_answers": "yes"}, "")["info"]
        assert info["eeoConsent"] is False        # explicit decline honored
        assert info["eeoSaveAnswers"] is True
        assert info["eeoShareAnswers"] is False

    def test_indeed_consent_defaults(self):
        info = onboard.build_onboarding_json({}, "")["info"]
        assert info["eeoConsent"] is True         # auto-agree the required consent
        assert info["eeoSaveAnswers"] is False and info["eeoShareAnswers"] is False


class TestParseResumeInfo:
    SAMPLE = (
        "Thomas Thirlwall\n"
        "+1 (956) 525-3015 | ththirlwall99@gmail.com | Dallas, TX | "
        "linkedin.com/in/thomas-thirlwall | https://github.com/FrameAutomata | "
        "https://thomas.dev\n"
        "PROFESSIONAL SUMMARY\nEngineer.\n"
    )

    def test_extracts_all_fields(self):
        info = onboard.parse_resume_info(self.SAMPLE)
        assert info["name"] == "Thomas Thirlwall"
        assert info["email"] == "ththirlwall99@gmail.com"
        assert info["phone"] == "+1 (956) 525-3015"
        assert info["location"] == "Dallas, TX"
        assert info["linkedin"] == "linkedin.com/in/thomas-thirlwall"
        assert info["github"] == "github.com/FrameAutomata"
        # website = first non-social URL
        assert info["website"] == "https://thomas.dev"

    def test_website_skips_social_urls(self):
        info = onboard.parse_resume_info("Jane\nhttps://github.com/jane only")
        assert info["website"] is None

    def test_missing_fields_are_none(self):
        info = onboard.parse_resume_info("Just Some Text\nno contacts here")
        assert info["email"] is None and info["phone"] is None
        assert info["name"] == "Just Some Text"

    def test_first_line_with_at_is_not_treated_as_name(self):
        info = onboard.parse_resume_info("me@example.com\nReal Name")
        assert info["name"] is None  # first line has '@'
        assert info["email"] == "me@example.com"


class TestCollectSecretBlobs:
    def test_reads_and_base64s_present_files(self, tmp_path):
        (tmp_path / "config").mkdir()
        (tmp_path / "resumes").mkdir()
        (tmp_path / "career-ops" / "config").mkdir(parents=True)
        (tmp_path / "career-ops" / "modes").mkdir(parents=True)
        (tmp_path / "config" / "search.yml").write_text("searches: []", encoding="utf-8")
        (tmp_path / "resumes" / "resume.txt").write_text("resume", encoding="utf-8")
        (tmp_path / "career-ops" / "cv.md").write_text("# CV", encoding="utf-8")
        (tmp_path / "career-ops" / "config" / "profile.yml").write_text("candidate: {}", encoding="utf-8")
        # PROFILE_MD intentionally absent — it's optional.
        blobs = onboard.collect_secret_blobs(tmp_path)
        assert set(blobs) >= set(onboard.REQUIRED_SECRETS)
        assert "PROFILE_MD_B64" not in blobs
        assert base64.b64decode(blobs["CV_MD_B64"]).decode() == "# CV"

    def test_raises_when_required_missing(self, tmp_path):
        with pytest.raises(onboard.OnboardError, match="search.yml"):
            onboard.collect_secret_blobs(tmp_path)


@pytest.mark.skipif(not _node_deps_available(),
                    reason="node or its npm deps (yaml/pdf-parse) not installed")
class TestNodeRoundTrip:
    """Guards against the interactive and --from-json paths drifting: run the
    real generator on a fixture and assert the four artifacts appear."""

    def test_from_json_produces_artifacts(self, tmp_path):
        repo = Path(__file__).resolve().parent.parent
        # Run node in an isolated cwd with only the example config available, so
        # we don't touch the real config/career-ops.
        work = tmp_path
        (work / "config").mkdir()
        shutil.copy(repo / "config" / "search.example.yml", work / "config" / "search.example.yml")
        (work / "career-ops" / "config").mkdir(parents=True)
        (work / "career-ops" / "modes").mkdir(parents=True)

        payload = onboard.build_onboarding_json(
            {"name": "Jane Dev", "target_roles": "Backend Engineer",
             "locations": "US Remote, Dallas, TX", "sites": ["indeed"],
             "citizenship": "Canadian", "requires_sponsorship": "yes",
             "work_auth_regions": "Canada", "eligible_countries": "Canada, United States",
             "work_permit_type": "Needs sponsorship"},
            resume_text="Jane Dev\nSKILLS\nPython, AWS\nEXPERIENCE\nAcme",
        )
        result = onboard.run_generation(work, payload)
        assert result.get("ok") is True
        profile_path = work / "career-ops" / "config" / "profile.yml"
        assert profile_path.exists()
        assert (work / "career-ops" / "cv.md").exists()
        assert (work / "career-ops" / "modes" / "_profile.md").exists()
        assert (work / "config" / "search.yml").exists()

        # Work-authorization answers flow form -> JSON -> generated profile.yml.
        import yaml
        profile = yaml.safe_load(profile_path.read_text(encoding="utf-8"))
        wa = profile["work_authorization"]
        assert wa["citizenship"] == "Canadian"
        assert wa["requires_sponsorship"] is True
        assert wa["legally_authorized_to_work_in"] == ["Canada"]
        assert wa["eligible_countries"] == ["Canada", "United States"]
        assert wa["work_permit_type"] == "Needs sponsorship"

    def test_derive_form_reads_back_what_the_generator_wrote(self, tmp_path):
        """The drift guard for derive_form.

        derive_form is a hand-written INVERSE of setup-profile.mjs's YAML
        authorship, and nothing else ties the two together — `generateProfile`
        renaming `candidate.portfolio_url` would leave the whole suite green
        while prefill silently returned nothing for that field. Which is
        invisible: "the box is empty" is the #145 symptom, not an error.

        A unit fixture can't catch that (it would be a third hand-written copy
        of the same shape). Running the REAL generator and reading its output
        back can. Asserted over the keys the input actually set — Node fills
        defaults the form left blank, and those are not drift.
        """
        repo = Path(__file__).resolve().parent.parent
        work = tmp_path
        (work / "config").mkdir()
        shutil.copy(repo / "config" / "search.example.yml", work / "config" / "search.example.yml")
        (work / "career-ops" / "config").mkdir(parents=True)
        (work / "career-ops" / "modes").mkdir(parents=True)

        form = {
            "name": "Jane Dev", "email": "jane@example.com", "phone": "+1 (555) 123-4567",
            "location": "Dallas, TX", "linkedin": "linkedin.com/in/janedev",
            "github": "github.com/janedev", "website": "janedev.dev",
            "street": "1 Main St", "state": "TX", "postal_code": "75201",
            "tailoring_instructions": "Lead with impact.",
            "target_roles": "Backend Engineer, Platform Engineer",
            "negative_roles": "Intern, Director",
            "comp_target": "$150K-190K", "comp_min": "$130K",
            "location_flexibility": "Remote preferred",
            "citizenship": "Canadian", "requires_sponsorship": "yes",
            "work_auth_regions": "Canada", "eligible_countries": "Canada, United States",
            "work_permit_type": "Needs sponsorship",
            "eeo_gender": "Female", "eeo_race": "Asian",
            "data_processing_consent": "no", "save_answers": "yes", "share_answers": "no",
            "locations": "US Remote, Dallas, TX", "sites": ["indeed"],
            "hours_old": "48", "results_wanted": "75", "distance": "25",
            "include_easy_apply": True,
        }
        onboard.run_generation(work, onboard.build_onboarding_json(form, "Jane Dev\nSKILLS\nPython"))

        derived = onboard.derive_form(work, work / "career-ops")
        for key, sent in form.items():
            assert key in derived, f"{key} was written but derive_form doesn't read it back"
            assert derived[key] == sent, f"{key} did not survive the round trip"


class TestExtractResumeText:
    """onboard.extract_resume_text(bytes, filename) dispatches by the uploaded
    filename's suffix so the UI accepts DOCX/ODT as well as PDF."""

    def test_extracts_docx_bytes(self, tmp_path):
        from docx import Document
        d = Document()
        d.add_paragraph("Jane Dev")
        d.add_paragraph("Python, AWS")
        f = tmp_path / "src.docx"
        d.save(str(f))
        text = onboard.extract_resume_text(f.read_bytes(), "Jane_Resume.docx")
        assert "Jane Dev" in text and "Python, AWS" in text

    def test_unsupported_format_raises(self):
        with pytest.raises(ValueError):
            onboard.extract_resume_text(b"x", "resume.rtf")


class TestRunGenerationErrors:
    def test_node_missing_raises_clear_error(self, tmp_path, mocker):
        mocker.patch("pipeline.app.onboard.subprocess.run", side_effect=FileNotFoundError())
        with pytest.raises(onboard.OnboardError, match="node not found"):
            onboard.run_generation(tmp_path, {"info": {}})

    def test_nonzero_exit_surfaces_stderr(self, tmp_path, mocker):
        cp = subprocess.CompletedProcess(args=["node"], returncode=1, stdout="", stderr="kaboom")
        mocker.patch("pipeline.app.onboard.subprocess.run", return_value=cp)
        with pytest.raises(onboard.OnboardError, match="kaboom"):
            onboard.run_generation(tmp_path, {"info": {}})


class TestSidecar:
    """The sidecar lets the wizard prefill every field on a revisit so the user
    only touches the knob they want to change. api_key must never land in it —
    the sidecar lives on disk; the key belongs in GitHub Secrets only."""

    def test_load_returns_none_when_missing(self, tmp_path):
        assert onboard.load_sidecar(tmp_path) is None

    def test_save_then_load_round_trips(self, tmp_path):
        payload = {"name": "Jane", "results_wanted": 5, "sites": ["indeed", "linkedin"]}
        onboard.save_sidecar(tmp_path, payload)
        assert onboard.load_sidecar(tmp_path) == payload

    def test_save_strips_api_key(self, tmp_path):
        onboard.save_sidecar(tmp_path, {"name": "Jane", "api_key": "secret-xyz"})
        loaded = onboard.load_sidecar(tmp_path)
        assert "api_key" not in loaded
        # Belt-and-braces: the raw file shouldn't contain the key either.
        raw = (tmp_path / ".ui-cache" / "onboarding.json").read_text(encoding="utf-8")
        assert "secret-xyz" not in raw

    def test_load_returns_none_on_corrupt_file(self, tmp_path):
        # If someone hand-edits the sidecar into invalid JSON, fall back to
        # "no prefill" rather than crash the wizard.
        sidecar = tmp_path / ".ui-cache" / "onboarding.json"
        sidecar.parent.mkdir(parents=True)
        sidecar.write_text("{ not valid json", encoding="utf-8")
        assert onboard.load_sidecar(tmp_path) is None

    def test_save_swallows_io_errors(self, tmp_path, mocker):
        # Sidecar persistence is a UX nicety. If we can't write (read-only fs,
        # antivirus lock, whatever), don't fail the onboarding submission —
        # the secrets have already been written at that point.
        mocker.patch("pipeline.app.onboard.Path.write_text",
                     side_effect=OSError("read-only"))
        # No exception:
        onboard.save_sidecar(tmp_path, {"name": "Jane"})


class TestConfiguredState:
    """`is_configured` answers "has this copy been set up", which the wizard used
    to ask as "is there a sidecar" — i.e. "did someone submit THIS WIZARD" —
    leaving every other setup route looking first-time forever (#145)."""

    def test_bare_checkout_is_not_configured(self, tmp_path):
        assert onboard.is_configured(tmp_path, tmp_path / "career-ops") is False

    def test_search_yml_alone_is_not_evidence(self, tmp_path):
        # setup.sh copies search.example.yml here before the user has answered
        # anything, so its existence says only that setup ran.
        (tmp_path / "config").mkdir()
        (tmp_path / "config" / "search.yml").write_text("searches: []", encoding="utf-8")
        assert onboard.is_configured(tmp_path, tmp_path / "career-ops") is False

    @pytest.mark.parametrize("rel", ["config/profile.yml", "cv.md"])
    def test_cli_generated_artifacts_count(self, tmp_path, rel):
        co = tmp_path / "career-ops"
        target = co / rel
        target.parent.mkdir(parents=True)
        target.write_text("x", encoding="utf-8")
        assert onboard.is_configured(tmp_path, co) is True

    def test_sidecar_still_counts(self, tmp_path):
        onboard.save_sidecar(tmp_path, {"name": "Jane"})
        assert onboard.is_configured(tmp_path, tmp_path / "career-ops") is True


class TestResumeOnFile:
    """The step-0 gate asks the disk, in the order onboard_submit falls back
    through — so the gate and the submit behind it can't disagree."""

    def test_absent_when_nothing_on_disk(self, tmp_path):
        assert onboard.resume_on_file(tmp_path) is False

    def test_probes_the_import_formats(self, tmp_path):
        (tmp_path / "resumes").mkdir()
        (tmp_path / "resumes" / "resume.docx").write_bytes(b"x")
        assert onboard.resume_on_file(tmp_path) is True

    def test_honours_resume_path(self, tmp_path, monkeypatch):
        # .env.example documents RESUME_PATH and extract_resume_text() honours
        # it, so a resume under any other name is present as far as the pipeline
        # is concerned. Matching the fixed names alone read it as absent.
        odd = tmp_path / "docs" / "my-cv-2026.docx"
        odd.parent.mkdir(parents=True)
        odd.write_bytes(b"x")
        monkeypatch.setenv("RESUME_PATH", "docs/my-cv-2026.docx")
        assert onboard.resume_on_file(tmp_path) is True

    def test_txt_sidecar_counts(self, tmp_path, monkeypatch):
        # What a previous submit left; onboard_submit reuses it verbatim.
        monkeypatch.delenv("RESUME_PATH", raising=False)
        (tmp_path / "resumes").mkdir()
        (tmp_path / "resumes" / "resume.txt").write_text("Jane Dev", encoding="utf-8")
        assert onboard.resume_on_file(tmp_path) is True


PROFILE_YML = """
candidate:
  full_name: Jane Dev
  email: jane@example.com
  phone: ""
  location: Dallas, TX
  linkedin: linkedin.com/in/janedev
  github: ""
  portfolio_url: janedev.dev
target_roles:
  primary:
    - Senior Backend Engineer
    - Platform Engineer
  archetypes:
    - {name: Senior Backend Engineer, level: senior, fit: primary}
    - {name: Platform Engineer, level: senior, fit: secondary}
    - {name: Staff Engineer, level: senior, fit: secondary}
compensation:
  target_range: $150K-190K
  minimum: $130K
  location_flexibility: Remote preferred
location:
  street: 1 Main St
  state: TX
  postal_code: "75201"
tailoring:
  instructions: Lead with impact.
work_authorization:
  citizenship: US
  legally_authorized_to_work_in: [United States, Canada]
  requires_sponsorship: false
  work_permit_type: Citizen
  eligible_countries: [United States]
voluntary_disclosures:
  gender: ""
  race_ethnicity: Asian
  veteran_status: ""
  disability_status: ""
  data_processing_consent: true
  save_answers: false
  share_answers: false
"""


@pytest.fixture
def profile():
    """PROFILE_YML parsed — a fresh dict per test, since some mutate it."""
    return yaml.safe_load(PROFILE_YML)


@pytest.fixture
def search():
    """SEARCH_YML parsed — fresh per test, as above."""
    return yaml.safe_load(SEARCH_YML)


class TestDeriveFromProfile:
    """Prefill reads the files a RUN uses, so a copy configured any way at all
    shows what is currently in effect before the wizard overwrites it."""

    @pytest.fixture
    def derived(self, profile):
        return onboard.derive_from_profile(profile)

    def test_maps_contact_and_address_fields(self, derived):
        assert derived["name"] == "Jane Dev"
        assert derived["email"] == "jane@example.com"
        assert derived["location"] == "Dallas, TX"
        assert derived["linkedin"] == "linkedin.com/in/janedev"
        assert derived["website"] == "janedev.dev"
        assert derived["street"] == "1 Main St"
        assert derived["postal_code"] == "75201"
        assert derived["tailoring_instructions"] == "Lead with impact."

    def test_blank_file_values_are_omitted(self, derived):
        # A blank in the file is a field the setup route never filled, not an
        # answer. Emitting "" would win the merge below and render as an empty
        # box — the exact symptom this exists to fix.
        assert "phone" not in derived
        assert "github" not in derived
        assert "eeo_gender" not in derived

    def test_target_roles_come_from_archetypes_not_primary(self, derived):
        # `primary` is the same list truncated to two, so reading it would drop
        # roles 3+ on the next save.
        assert derived["target_roles"] == "Senior Backend Engineer, Platform Engineer, Staff Engineer"

    def test_falls_back_to_primary_without_archetypes(self, profile):
        del profile["target_roles"]["archetypes"]
        assert onboard.derive_from_profile(profile)["target_roles"] == (
            "Senior Backend Engineer, Platform Engineer")

    def test_lists_render_as_the_csv_the_form_uses(self, derived):
        assert derived["work_auth_regions"] == "United States, Canada"
        assert derived["eligible_countries"] == "United States"

    def test_booleans_keep_both_answers(self, derived):
        # Keyed on the key being present, not the value being non-blank: `false`
        # is an answer, and it serializes the way the form's select/checkbox do.
        assert derived["requires_sponsorship"] == "no"
        assert derived["data_processing_consent"] == "yes"
        assert derived["save_answers"] == "no"

    def test_missing_booleans_are_omitted(self):
        assert "requires_sponsorship" not in onboard.derive_from_profile({})

    @pytest.mark.parametrize("written,expected", [
        (False, "no"), ("false", "no"), ("no", "no"), ("No", "no"),
        (True, "yes"), ("true", "yes"), ("yes", "yes"),
    ])
    def test_quoted_yes_no_is_not_read_as_truthiness(self, written, expected):
        # YAML leaves `requires_sponsorship: "false"` a truthy Python string, so
        # plain truthiness would prefill "yes" — and the next Save would write
        # that inversion back into profile.yml. Hand-edited copies are exactly
        # the population this feature is for.
        derived = onboard.derive_from_profile(
            {"work_authorization": {"requires_sponsorship": written}})
        assert derived["requires_sponsorship"] == expected

    def test_unreadable_boolean_falls_through_to_the_sidecar(self):
        # "maybe" is not an answer; guessing either way would be worse than
        # letting the last submitted value stand.
        derived = onboard.derive_from_profile(
            {"work_authorization": {"requires_sponsorship": "maybe"}})
        assert "requires_sponsorship" not in derived

    def test_garbage_shapes_do_not_raise(self):
        # A half-written profile.yml must degrade to "nothing to prefill".
        assert onboard.derive_from_profile({"candidate": "not a mapping"}) == {}
        assert onboard.derive_from_profile({}) == {}


SEARCH_YML = """
searches:
  - name: US Remote
    sites: [indeed, linkedin]
    results_wanted: 100
    location: United States
    country_indeed: USA
    is_remote: true
  - name: Dallas, TX
    sites: [linkedin]
    results_wanted: 100
    location: Dallas, TX
    country_indeed: USA
    hours_old: 24
    distance: 25
  - name: easy apply
    sites: [indeed]
    results_wanted: 100
    location: Dallas, TX
    country_indeed: USA
    easy_apply: "true"
filter:
  negative_titles: [intern, director]
"""


class TestDeriveFromSearch:
    @pytest.fixture
    def derived(self, search):
        return onboard.derive_from_search(search)

    def test_locations_round_trip_through_the_form(self, derived):
        # The wizard's `locations` field is re-read by parse_locations and
        # rebuilt by build_onboarding_json, so what we derive has to survive
        # that trip and land on the same passes.
        assert derived["locations"] == "US Remote, Dallas, TX"
        entries = onboard.build_onboarding_json(
            {**derived, "distance": derived["distance"]}, "text",
        )["searchSettings"]["locations"]
        assert entries[0] == {"raw": "US Remote", "isRemote": True,
                              "country": "USA", "location": "United States"}
        assert entries[1] == {"raw": "Dallas, TX", "isRemote": False,
                              "country": "USA", "location": "Dallas, TX",
                              "distance": 25}

    def test_pass_name_is_preferred_when_it_still_describes_the_pass(self, derived):
        # setup-profile.mjs writes `name: loc.raw` — the text the user typed. A
        # wizard user who typed "US Remote" must not have it rewritten to the
        # equivalent "Remote United States"; the banner promises the fields show
        # what's in effect, not a restatement of it.
        assert derived["locations"].startswith("US Remote")

    def test_stale_pass_name_loses_to_what_the_pass_searches(self):
        # A hand-edited config whose name no longer matches: the name is a label
        # ("recent local" in the shipped example), not a location, so rebuilding
        # from the fields is the only honest answer.
        derived = onboard.derive_from_search({"searches": [
            {"name": "recent local", "location": "Dallas, TX", "hours_old": 24}]})
        assert derived["locations"] == "Dallas, TX"

    def test_remote_prefix_survives_a_comma_bearing_location(self):
        # "Remote (Dallas, TX)" would split — parse_locations rejoins "City, ST"
        # pairs and "TX)" is not a state code — so the prefix spelling is the
        # one that holds when there's no usable name to fall back on.
        derived = onboard.derive_from_search(
            {"searches": [{"location": "Dallas, TX", "is_remote": True}]})
        assert onboard.parse_locations(derived["locations"]) == ["Remote Dallas, TX"]

    def test_easy_apply_pass_is_a_flag_not_a_location(self, derived):
        assert derived["include_easy_apply"] is True
        assert "easy apply" not in derived["locations"]

    def test_easy_apply_false_when_no_such_pass(self):
        derived = onboard.derive_from_search(
            {"searches": [{"location": "Dallas, TX", "hours_old": 24}]})
        assert derived["include_easy_apply"] is False

    def test_easy_apply_read_the_way_jobspy_reads_it(self):
        # normalize_pass first: `easy_apply: false` sends no filter, so it is
        # not an easy-apply pass however present the key is.
        derived = onboard.derive_from_search(
            {"searches": [{"location": "Dallas, TX", "easy_apply": False}]})
        assert derived["include_easy_apply"] is False

    def test_numbers_and_boards_and_negative_titles(self, derived):
        assert derived["results_wanted"] == "100"
        assert derived["hours_old"] == "24"
        assert derived["distance"] == "25"
        assert derived["sites"] == ["indeed", "linkedin"]
        assert derived["negative_roles"] == "intern, director"

    def test_missing_sites_inherits_the_supported_boards(self):
        # resolve_sites' rule, not a second reading of it: an omitted `sites`
        # scrapes every supported board.
        derived = onboard.derive_from_search({"searches": [{"location": "Dallas, TX"}]})
        assert derived["sites"] == list(SUPPORTED_SITES)

    def test_legacy_single_search_mapping(self):
        derived = onboard.derive_from_search(
            {"search": {"location": "Dallas, TX", "results_wanted": 40}})
        assert derived["locations"] == "Dallas, TX"
        assert derived["results_wanted"] == "40"

    def test_no_passes_yields_no_search_fields(self):
        derived = onboard.derive_from_search({"filter": {"negative_titles": ["intern"]}})
        assert derived == {"negative_roles": "intern"}


class TestSearchDetailAtRisk:
    """Saving rewrites `searches:` wholesale from six form fields. That cost
    nothing while only wizard-written configs reached the wizard; hand-written
    ones now do, so what a Save would take away has to be said out loud."""

    def _cfg(self, tmp_path, text):
        p = tmp_path / "search.yml"
        p.write_text(text, encoding="utf-8")
        return p

    def test_silent_for_a_wizard_written_config(self, tmp_path):
        # The warning has to be precise or it is wallpaper: every field here is
        # one the wizard writes and reads back.
        assert onboard.search_detail_at_risk(self._cfg(tmp_path, SEARCH_YML)) == []

    def test_silent_for_the_shipped_example(self, tmp_path):
        example = Path(__file__).resolve().parent.parent / "config" / "search.example.yml"
        assert onboard.search_detail_at_risk(example) == []

    def test_names_pass_keys_the_wizard_cannot_express(self, tmp_path):
        risk = onboard.search_detail_at_risk(self._cfg(tmp_path, """
searches:
  - name: Dallas
    location: "Dallas, TX"
    hours_old: 24
    job_type: fulltime
    linkedin_company_ids: [1441]
    offset: 25
"""))
        assert "job_type" in risk and "linkedin_company_ids" in risk and "offset" in risk

    def test_names_a_location_the_locations_field_would_split(self, tmp_path):
        # parse_locations rejoins only "City, ST" pairs, so this one pass comes
        # back as two ("Berlin" and "Germany"), each re-targeted to the USA.
        risk = onboard.search_detail_at_risk(self._cfg(tmp_path, """
searches:
  - {name: Berlin, location: "Berlin, Germany", country_indeed: Germany, hours_old: 24}
"""))
        assert any("Berlin, Germany" in r for r in risk)

    def test_names_a_country_the_wizard_would_re_infer_differently(self, tmp_path):
        risk = onboard.search_detail_at_risk(self._cfg(tmp_path, """
searches:
  - {name: Springfield, location: Springfield, country_indeed: Canada, hours_old: 24}
"""))
        assert any("country_indeed" in r for r in risk)

    def test_missing_config_is_not_a_warning(self, tmp_path):
        assert onboard.search_detail_at_risk(tmp_path / "nope.yml") == []


class TestPrefillMerge:
    """Real files win where they have a home; the sidecar keeps the rest."""

    def _write(self, tmp_path):
        co = tmp_path / "career-ops"
        (co / "config").mkdir(parents=True)
        (co / "config" / "profile.yml").write_text(PROFILE_YML, encoding="utf-8")
        (tmp_path / "config").mkdir()
        (tmp_path / "config" / "search.yml").write_text(SEARCH_YML, encoding="utf-8")
        return co

    def test_files_win_over_a_stale_sidecar(self, tmp_path):
        co = self._write(tmp_path)
        onboard.save_sidecar(tmp_path, {"name": "Old Name", "results_wanted": "5"})
        merged = onboard.prefill_form(tmp_path, co)
        assert merged["name"] == "Jane Dev"
        assert merged["results_wanted"] == "100"

    def test_sidecar_supplies_what_the_files_have_no_home_for(self, tmp_path):
        # The Narrative step lands in _profile.md as prose under headings the
        # user is invited to rewrite, so it is read back from the sidecar alone.
        co = self._write(tmp_path)
        onboard.save_sidecar(tmp_path, {"exit_story": "Moving to platform work.",
                                        "deal_breakers": "On-site 5 days"})
        merged = onboard.prefill_form(tmp_path, co)
        assert merged["exit_story"] == "Moving to platform work."
        assert merged["deal_breakers"] == "On-site 5 days"

    def test_cli_configured_copy_prefills_with_no_sidecar_at_all(self, tmp_path):
        merged = onboard.prefill_form(tmp_path, self._write(tmp_path))
        assert merged["name"] == "Jane Dev"
        assert merged["locations"] == "US Remote, Dallas, TX"

    def test_bare_checkout_prefills_nothing(self, tmp_path):
        assert onboard.prefill_form(tmp_path, tmp_path / "career-ops") == {}

    def test_every_derived_key_names_a_real_form_field(self, html, profile, search):
        # prefillForm looks each key up as `[name="<key>"]`, so a derived key
        # that doesn't name a field is prefill that silently does nothing —
        # which is the symptom (blank fields), not an error anyone would see.
        # Guards a rename on either side, over EVERY key the two derivations can
        # emit (the blank-omitting rule would otherwise hide the EEO fields).
        profile["voluntary_disclosures"].update(
            gender="Female", veteran_status="No", disability_status="No")
        keys = set(onboard.derive_from_profile(profile))
        keys |= set(onboard.derive_from_search(search))
        assert len(keys) > 20, "fixtures should exercise most of the form"
        for key in keys:
            assert re.search(rf'name="{re.escape(key)}"', html), \
                f"derive_form emits {key!r}, which no onboard.html field is named"

    def test_unreadable_yaml_degrades_to_the_sidecar(self, tmp_path):
        co = self._write(tmp_path)
        (co / "config" / "profile.yml").write_text("{ not: [valid", encoding="utf-8")
        onboard.save_sidecar(tmp_path, {"name": "Jane"})
        merged = onboard.prefill_form(tmp_path, co)
        assert merged["name"] == "Jane"
        assert merged["locations"] == "US Remote, Dallas, TX"


class TestArticleDigestSecretWiring:
    """article-digest.md is generated during onboarding (LLM-grounded) and must
    ride along to the cloud as an OPTIONAL secret — present when generation
    succeeded, silently omitted when it didn't (so onboarding never hard-depends
    on the digest)."""

    def _seed_required(self, root):
        (root / "config").mkdir()
        (root / "resumes").mkdir()
        (root / "career-ops" / "config").mkdir(parents=True)
        (root / "career-ops" / "modes").mkdir(parents=True)
        (root / "config" / "search.yml").write_text("searches: []", encoding="utf-8")
        (root / "resumes" / "resume.txt").write_text("resume", encoding="utf-8")
        (root / "career-ops" / "cv.md").write_text("# CV", encoding="utf-8")
        (root / "career-ops" / "config" / "profile.yml").write_text("candidate: {}", encoding="utf-8")

    def test_article_digest_is_an_optional_secret(self):
        assert onboard.SECRET_FILES.get("ARTICLE_DIGEST_B64") == "career-ops/article-digest.md"
        assert "ARTICLE_DIGEST_B64" not in onboard.REQUIRED_SECRETS

    def test_collect_includes_article_digest_when_present(self, tmp_path):
        self._seed_required(tmp_path)
        (tmp_path / "career-ops" / "article-digest.md").write_text("## Proj", encoding="utf-8")
        blobs = onboard.collect_secret_blobs(tmp_path)
        assert "ARTICLE_DIGEST_B64" in blobs
        assert base64.b64decode(blobs["ARTICLE_DIGEST_B64"]).decode() == "## Proj"

    def test_collect_omits_article_digest_when_absent(self, tmp_path):
        self._seed_required(tmp_path)
        blobs = onboard.collect_secret_blobs(tmp_path)
        assert "ARTICLE_DIGEST_B64" not in blobs


class TestProfileMasterSecretWiring:
    """The living PROFILE.md (grown by the browser agent in HANDOFF_OUT_DIR) rides
    to the cloud as an OPTIONAL PROFILE_MASTER_B64 secret, so the cloud evaluator
    scores against the same master as local eval. Special-cased (not in
    SECRET_FILES) because it lives at HANDOFF_OUT_DIR, not a repo-relative path."""

    def _seed_required(self, root):
        (root / "config").mkdir()
        (root / "resumes").mkdir()
        (root / "career-ops" / "config").mkdir(parents=True)
        (root / "config" / "search.yml").write_text("searches: []", encoding="utf-8")
        (root / "resumes" / "resume.txt").write_text("resume", encoding="utf-8")
        (root / "career-ops" / "cv.md").write_text("# CV", encoding="utf-8")
        (root / "career-ops" / "config" / "profile.yml").write_text("candidate: {}", encoding="utf-8")

    def test_profile_master_is_not_required(self):
        assert "PROFILE_MASTER_B64" not in onboard.REQUIRED_SECRETS

    def test_collect_includes_profile_master_when_present(self, tmp_path, monkeypatch):
        self._seed_required(tmp_path)
        handoff_dir = tmp_path / "handoff"
        handoff_dir.mkdir()
        (handoff_dir / "PROFILE.md").write_text("LIVING MASTER PROFILE", encoding="utf-8")
        monkeypatch.setenv("HANDOFF_OUT_DIR", str(handoff_dir))
        monkeypatch.setenv("CAREER_OPS_PATH", str(tmp_path / "career-ops"))   # pin the fallback
        blobs = onboard.collect_secret_blobs(tmp_path)
        assert "PROFILE_MASTER_B64" in blobs
        assert base64.b64decode(blobs["PROFILE_MASTER_B64"]).decode() == "LIVING MASTER PROFILE"

    def test_collect_omits_profile_master_when_absent(self, tmp_path, monkeypatch):
        self._seed_required(tmp_path)
        handoff_dir = tmp_path / "handoff"
        handoff_dir.mkdir()                          # no PROFILE.md anywhere
        monkeypatch.setenv("HANDOFF_OUT_DIR", str(handoff_dir))
        monkeypatch.setenv("CAREER_OPS_PATH", str(tmp_path / "career-ops"))   # fallback dir has none either
        blobs = onboard.collect_secret_blobs(tmp_path)
        assert "PROFILE_MASTER_B64" not in blobs

    def test_collect_skips_oversized_profile_master(self, tmp_path, monkeypatch, capsys):
        # An append-only PROFILE.md that outgrows GitHub's secret cap must degrade
        # to the seeds (skip + log), not 502 the onboard / strand the provider key.
        self._seed_required(tmp_path)
        handoff_dir = tmp_path / "handoff"
        handoff_dir.mkdir()
        big = "x" * onboard.PROFILE_MASTER_MAX_B64   # base64 ~1.33x → over the cap
        (handoff_dir / "PROFILE.md").write_text(big, encoding="utf-8")
        monkeypatch.setenv("HANDOFF_OUT_DIR", str(handoff_dir))
        monkeypatch.setenv("CAREER_OPS_PATH", str(tmp_path / "career-ops"))
        blobs = onboard.collect_secret_blobs(tmp_path)
        assert "PROFILE_MASTER_B64" not in blobs      # skipped, not a crash
        assert "too large" in capsys.readouterr().out  # and it says why


class TestEeoDatalists:
    """The four voluntary self-ID fields offer standard answers as dropdown
    suggestions, but must stay free-text so a user can type the employer's exact
    wording (the answers are fuzzy-matched per form) or a non-US category. That
    means an <input list="..."> wired to a <datalist> — never a hard <select>,
    which would lock in one (US-centric) option set and drop free typing."""

    EEO_FIELDS = ["eeo_gender", "eeo_race", "eeo_veteran", "eeo_disability"]

    def _input_tag(self, html, name):
        m = re.search(rf'<input\b[^>]*\bname="{re.escape(name)}"[^>]*>', html)
        assert m, f"no <input name={name!r}> found"
        return m.group(0)

    def test_each_eeo_field_is_input_not_select(self, html):
        # A <select name="eeo_*"> would mean we lost free-text entry.
        for name in self.EEO_FIELDS:
            assert not re.search(rf'<select\b[^>]*\bname="{re.escape(name)}"', html), \
                f"{name} must stay an <input> (free-text), not a <select>"
            self._input_tag(html, name)  # the <input> still exists

    def test_each_eeo_input_references_an_existing_datalist(self, html):
        datalist_ids = set(re.findall(r'<datalist\b[^>]*\bid="([^"]+)"', html))
        for name in self.EEO_FIELDS:
            tag = self._input_tag(html, name)
            m = re.search(r'\blist="([^"]+)"', tag)
            assert m, f"{name} input is missing a list= attribute"
            assert m.group(1) in datalist_ids, \
                f"{name} list={m.group(1)!r} has no matching <datalist id>"

    def test_each_eeo_input_reenables_autocomplete(self, html):
        # The form sets autocomplete="off" (so the browser doesn't autofill
        # name/email across the wizard). Chromium suppresses <datalist>
        # suggestions whenever the input's effective autocomplete is "off", so
        # each EEO input must override it back to "on" or the dropdown is dead.
        for name in self.EEO_FIELDS:
            tag = self._input_tag(html, name)
            m = re.search(r'\bautocomplete="([^"]+)"', tag)
            assert m, f"{name} input must set autocomplete to override the form default"
            assert m.group(1).lower() != "off", \
                f"{name} input has autocomplete=off — its datalist won't show in Chromium"

    def test_datalists_offer_standard_and_decline_options(self, html):
        # Spot-check that the suggestion sets carry the usual wording, including
        # an explicit "prefer not to say" so declining is one click, not just blank.
        assert "Non-binary" in html
        assert "Black or African American" in html
        assert "protected veteran" in html
        assert "do not have a disability" in html
        assert re.search(r"[Pp]refer not to (say|answer)|don't wish to answer|do not want to answer", html)


class TestPortfolioUrls:
    def test_combines_portfolio_csv_and_github(self):
        urls = onboard.portfolio_urls(
            {"portfolio": "https://a.com, https://b.com", "github": "https://github.com/me"})
        assert urls == ["https://a.com", "https://b.com", "https://github.com/me"]

    def test_accepts_list_and_skips_blank_github(self):
        assert onboard.portfolio_urls({"portfolio": ["https://a.com"], "github": ""}) == ["https://a.com"]

    def test_empty_when_nothing_supplied(self):
        assert onboard.portfolio_urls({}) == []


class TestOnboardHtmlSites:
    """The wizard must only offer the boards that actually produce rows.
    Glassdoor and ZipRecruiter are Cloudflare-403-walled and Google Jobs drops
    connections mid-response (crashing jobspy), so indeed + linkedin are the
    only supported checkboxes."""

    def test_offers_exactly_the_supported_boards(self, html):
        offered = set(re.findall(r'<input\b[^>]*\bname="sites"[^>]*\bvalue="([^"]+)"', html))
        assert offered == set(SUPPORTED_SITES)


class TestSupportedSitesMirror:
    """setup-profile.mjs restates SUPPORTED_SITES because Node can't import the
    Python constant. That mirror is what writes search.yml, so drift between the
    two silently reintroduces a dead board into every generated config."""

    def test_mjs_mirror_matches_the_python_constant(self):
        root = Path(__file__).resolve().parent.parent
        src = (root / "setup-profile.mjs").read_text(encoding="utf-8")
        m = re.search(r"const SUPPORTED_SITES = \[([^\]]*)\]", src)
        assert m, "no SUPPORTED_SITES literal found in setup-profile.mjs"
        mirrored = tuple(re.findall(r"'([^']+)'", m.group(1)))
        assert mirrored == SUPPORTED_SITES
