"""Tests for the deterministic résumé toolchain (Commit 3b).

`resume_render.render_docx` builds a one-page .docx from a generic content-JSON;
`resume_fit` gates its one-page fill %. Both are user-agnostic and LLM-free — the
LLM build/tailor loop that drives them to the 92–96% target is 3c. docx→PDF reuses
resume_tailor.render_pdf/page_count.
"""
import pytest

from pipeline import resume_build, resume_fit, resume_render


# A representative content-JSON: the generic schema (no personal/AI-specific keys).
_CONTENT = {
    "name": "Jane Doe",
    "contact": "jane@example.com · +1 (555) 010-2030 · Dallas, TX · linkedin.com/in/jane-doe",
    "summary": "Full-stack engineer who ships AI-first systems end to end, from schema to UI.",
    "skills": [
        {"label": "Languages", "items": "Python · Go · SQL"},
        {"label": "Cloud & CI/CD", "items": "AWS · Docker · GitHub Actions"},
    ],
    "experience": [
        {"org": "Acme Corp", "dates": "2022 – Present", "role": "Senior Engineer", "loc": "Remote",
         "bullets": ["Cut p99 latency 40% and saved $2M/yr", "Shipped billing to 5M users"]},
    ],
    "projects_heading": "Selected Projects",
    "projects": [
        {"org": "Job-Search Pipeline", "role": "Python · FastAPI · open source",
         "bullets": ["Orchestrated 7 LLM providers with automatic failover"]},
    ],
    "education": ["B.S. Computer Science — UTRGV", "AWS Certified Cloud Practitioner"],
}


def _paras(path):
    from docx import Document
    return [p.text for p in Document(str(path)).paragraphs]


class TestRenderDocx:
    """render_docx(content, out_path) → a .docx that carries every non-empty
    section, in the right order, with the flex section order-flippable."""

    def _render(self, tmp_path, content=None, name="r.docx"):
        return resume_render.render_docx(content or _CONTENT, tmp_path / name)

    def test_writes_docx_and_returns_path(self, tmp_path):
        out = self._render(tmp_path)
        assert out.exists() and out.suffix == ".docx"

    def test_identity_and_summary_present(self, tmp_path):
        t = "\n".join(_paras(self._render(tmp_path)))
        assert "Jane Doe" in t and "jane@example.com" in t
        assert "ships AI-first systems end to end" in t

    def test_skills_label_and_items(self, tmp_path):
        t = "\n".join(_paras(self._render(tmp_path)))
        assert "Languages" in t and "Python · Go · SQL" in t
        assert "AWS · Docker · GitHub Actions" in t

    def test_experience_fields_and_bullets(self, tmp_path):
        t = "\n".join(_paras(self._render(tmp_path)))
        for s in ("Acme Corp", "2022 – Present", "Senior Engineer", "Remote",
                  "Cut p99 latency 40% and saved $2M/yr", "Shipped billing to 5M users"):
            assert s in t

    def test_projects_and_education(self, tmp_path):
        t = "\n".join(_paras(self._render(tmp_path)))
        assert "Orchestrated 7 LLM providers with automatic failover" in t
        assert "B.S. Computer Science — UTRGV" in t
        assert "AWS Certified Cloud Practitioner" in t

    def test_all_section_headings_present(self, tmp_path):
        up = [p.upper() for p in _paras(self._render(tmp_path))]
        for h in ("PROFESSIONAL SUMMARY", "SKILLS", "PROFESSIONAL EXPERIENCE",
                  "SELECTED PROJECTS", "EDUCATION"):
            assert h in up

    def test_experience_before_projects_by_default(self, tmp_path):
        up = [p.upper() for p in _paras(self._render(tmp_path))]
        assert up.index("PROFESSIONAL EXPERIENCE") < up.index("SELECTED PROJECTS")

    def test_projects_first_hoists_the_flex_section(self, tmp_path):
        c = {**_CONTENT, "projects_first": True}
        up = [p.upper() for p in _paras(self._render(tmp_path, c))]
        assert up.index("SELECTED PROJECTS") < up.index("PROFESSIONAL EXPERIENCE")

    def test_empty_sections_omitted(self, tmp_path):
        c = {"name": "Jane Doe", "contact": "x", "summary": "s",
             "experience": [{"org": "A", "dates": "d", "role": "r", "loc": "l",
                             "bullets": ["b"]}]}
        up = [p.upper() for p in _paras(self._render(tmp_path, c))]
        assert "SKILLS" not in up
        assert "SELECTED PROJECTS" not in up and "PROJECTS" not in up
        assert "EDUCATION" not in up

    def test_projects_heading_defaults_when_absent(self, tmp_path):
        # A candidate who supplies `projects` but no `projects_heading` still gets
        # a labelled section (generic default "Projects").
        c = {k: v for k, v in _CONTENT.items() if k != "projects_heading"}
        up = [p.upper() for p in _paras(self._render(tmp_path, c))]
        assert "PROJECTS" in up

    def test_null_field_values_do_not_crash(self, tmp_path):
        # An LLM feeds the content-JSON and readily emits `null` for a field; render
        # must tolerate it (a `.get(k, "")` default only covers ABSENT keys).
        content = {"name": "Jane", "contact": "x", "summary": None,
                   "skills": [{"label": None, "items": None}],
                   "experience": [{"org": None, "dates": None, "role": None,
                                   "loc": None, "bullets": None}],
                   "projects": None, "education": [None]}
        assert resume_render.render_docx(content, tmp_path / "n.docx").exists()


class TestVerdict:
    """_verdict(pages, fill) — the pure fill-gate logic, no PDF needed."""

    def test_one_full_page_ok(self):
        assert resume_fit._verdict(1, 0.94) == (0, "OK")

    def test_two_pages_overfull(self):
        assert resume_fit._verdict(2, 0.50) == (3, "OVERFULL")

    def test_too_tight_flagged(self):
        assert resume_fit._verdict(1, 0.99) == (3, "TIGHT")

    @pytest.mark.parametrize("fill,verdict", [
        (0.879, "UNDERFULL"), (0.88, "OK"), (0.94, "OK"), (0.96, "OK"), (0.985, "OK")])
    def test_fill_boundaries(self, fill, verdict):
        assert resume_fit._verdict(1, fill)[1] == verdict

    def test_thresholds_are_the_agreed_band(self):
        assert (resume_fit.WARN_LO, resume_fit.TARGET_LO, resume_fit.TARGET_HI) == (0.88, 0.92, 0.96)


class TestCheckFit:
    """check_fit(pdf, content=None) — measure + verdict + content-lint. measure is
    monkeypatched so the gate logic is testable without LibreOffice."""

    def _patch(self, monkeypatch, pages, fill, blank_in=0.5):
        monkeypatch.setattr(resume_fit, "measure",
                            lambda p: resume_fit.Measurement(pages, fill, blank_in))

    def test_ok_result(self, monkeypatch):
        self._patch(monkeypatch, 1, 0.94)
        r = resume_fit.check_fit("x.pdf")
        assert r.ok and r.code == 0 and r.pages == 1 and r.verdict == "OK"

    def test_overfull_not_ok(self, monkeypatch):
        self._patch(monkeypatch, 2, 0.60)
        assert not resume_fit.check_fit("x.pdf").ok

    def test_underfull_hint_mentions_more_lines(self, monkeypatch):
        self._patch(monkeypatch, 1, 0.80, blank_in=1.5)
        r = resume_fit.check_fit("x.pdf")
        assert not r.ok and r.code == 2
        assert any("more line" in n.lower() for n in r.notes)

    def test_thin_flex_section_flagged(self, monkeypatch):
        # Fill is OK but the flex section is thin (1 entry / 1 bullet) → the
        # content-lint names it as the #1 fill lever.
        self._patch(monkeypatch, 1, 0.90)
        thin = {"projects_heading": "Selected Projects",
                "projects": [{"org": "P", "role": "r", "bullets": ["one"]}]}
        notes = resume_fit.check_fit("x.pdf", content=thin).notes
        assert any(("expand" in n.lower() or "thin" in n.lower()) for n in notes)

    def test_no_content_means_no_thin_flex_note(self, monkeypatch):
        self._patch(monkeypatch, 1, 0.94)
        r = resume_fit.check_fit("x.pdf")
        assert r.ok and not any("thin" in n.lower() for n in r.notes)

    def test_thin_flex_flagged_with_default_heading_and_order(self, monkeypatch):
        # The lint keys off the actual `projects` section, not the heading/order
        # flags — so a genuinely thin default-heading section is caught (it was a
        # silent false-negative when gated on the flags).
        self._patch(monkeypatch, 1, 0.90)
        thin = {"projects": [{"org": "P", "bullets": ["one"]}]}
        assert any("thin" in n.lower() for n in resume_fit.check_fit("x.pdf", thin).notes)

    def test_flag_without_projects_does_not_warn(self, monkeypatch):
        # projects_first set but NO projects → render emits no flex section, so the
        # lint must not warn about a section that isn't there (was a false-positive).
        self._patch(monkeypatch, 1, 0.94)
        r = resume_fit.check_fit("x.pdf", {"projects_first": True})
        assert not any("thin" in n.lower() for n in r.notes)


class TestScale:
    """Adaptive layout scale: render_docx(scale) resizes the WHOLE layout (font +
    spacing + line height) together, so the same content can be grown to fill a
    page or kept tight. scale=1.0 is the unchanged baseline."""

    def _name_pt(self, tmp_path, scale):
        from docx import Document
        p = resume_render.render_docx(_CONTENT, tmp_path / "s.docx", scale=scale)
        for para in Document(str(p)).paragraphs:
            if "Jane Doe" in para.text and para.runs:
                return para.runs[0].font.size.pt

    def test_default_scale_is_baseline(self, tmp_path):
        assert self._name_pt(tmp_path, 1.0) == pytest.approx(resume_render.NAME_PT)

    def test_scale_enlarges_fonts_proportionally(self, tmp_path):
        assert self._name_pt(tmp_path, 1.2) == pytest.approx(resume_render.NAME_PT * 1.2)


class TestSearchScale:
    """The pure scale search — the largest scale that still fits one page —
    exercised against a fake 'measure at scale', so it needs no LibreOffice."""

    def _fake(self, overflow_above):
        # One page at/below `overflow_above`, spills to two above it.
        return lambda s: resume_fit.Measurement(1 if s <= overflow_above else 2, 0.9, 0.5)

    def test_finds_the_overflow_boundary(self):
        s = resume_build._search_scale(self._fake(1.20), lo=0.95, hi=1.35, steps=8)
        assert 1.15 <= s <= 1.25

    def test_light_content_takes_the_max_scale(self):
        # Never overflows → the fullest one-pager is the max scale.
        assert resume_build._search_scale(self._fake(9.9), lo=0.95, hi=1.35, steps=6) \
            == pytest.approx(1.35)

    def test_heavy_content_falls_back_to_min_scale(self):
        # Overflows even at the smallest scale → min (the 3c loop must then trim).
        assert resume_build._search_scale(self._fake(0.0), lo=0.95, hi=1.35, steps=6) \
            == pytest.approx(0.95)

    def test_stops_at_fill_ceiling_before_page_spills(self):
        # Fill rises with scale and crosses TARGET_HI while still ONE page — the
        # binding constraint in production. The search must stop at the fill
        # ceiling, not run up to the page-spill boundary. (The _fake above holds
        # fill at 0.9, so this is the only case exercising the fill half of `ok`.)
        def fake(s):
            return resume_fit.Measurement(1, 0.80 + (s - 0.9) * 0.5, 0.5)   # 0.80 → 1.025
        s = resume_build._search_scale(fake, lo=0.9, hi=1.35, steps=12)
        assert fake(s).fill <= resume_fit.TARGET_HI
        assert fake(s).fill > resume_fit.TARGET_HI - 0.03            # and landed close to it


class TestFitToPage:
    """Integration: fit_to_page renders content at the fitted scale and returns the
    chosen PDF + scale + FitResult. Skipped without LibreOffice."""

    def _skip_if_no_soffice(self):
        from pipeline import resume_tailor
        if resume_tailor.find_soffice() is None:
            pytest.skip("LibreOffice (soffice) not installed")

    def test_light_content_scales_up_to_fill_one_page(self, tmp_path):
        self._skip_if_no_soffice()
        light = {"name": "Jane Doe", "contact": "jane@example.com · Dallas, TX",
                 "summary": "Short summary of a junior engineer.",
                 "experience": [{"org": "Acme", "dates": "2023 – Present", "role": "Engineer",
                                 "loc": "Remote", "bullets": ["Did a thing", "Did another thing"]}]}
        r = resume_build.fit_to_page(light, tmp_path)
        assert r.fit.pages == 1 and r.scale > 1.0 and r.pdf.exists()

    def test_result_carries_pdf_scale_and_fit(self, tmp_path):
        self._skip_if_no_soffice()
        r = resume_build.fit_to_page(_CONTENT, tmp_path)
        assert r.pdf.exists() and r.fit.pages == 1 and 0.9 <= r.scale <= 1.35


class TestRenderToPdfIntegration:
    """The full deterministic path render → PDF → measure, when LibreOffice is
    available (skipped otherwise, like the tailor's other soffice-bound tests)."""

    def test_render_measure_roundtrip(self, tmp_path):
        from pipeline import resume_tailor
        if resume_tailor.find_soffice() is None:
            pytest.skip("LibreOffice (soffice) not installed")
        docx = resume_render.render_docx(_CONTENT, tmp_path / "r.docx")
        pdf = resume_tailor.render_pdf(docx, tmp_path)
        assert pdf is not None and resume_tailor.page_count(pdf) == 1
        m = resume_fit.measure(pdf)
        assert m.pages == 1 and 0.0 < m.fill <= 1.0
