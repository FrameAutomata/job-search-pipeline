"""Tests for pipeline/filter.py"""

from datetime import datetime, timedelta
from pathlib import Path

import pandas as pd
import pytest

from pipeline import filter as filter_mod
from tests.conftest import SYNTHETIC_RESUME


class TestParseDatePosted:
    """Test filter.parse_date_posted function."""

    def test_parse_datetime_with_time_component(self):
        """Parse ISO datetime with time component."""
        result = filter_mod.parse_date_posted("2026-05-12 14:30:00")
        assert result == datetime(2026, 5, 12, 14, 30, 0)

    def test_parse_date_only(self):
        """Parse date only (no time)."""
        result = filter_mod.parse_date_posted("2026-05-12")
        assert result == datetime(2026, 5, 12, 0, 0, 0)

    def test_parse_empty_string_returns_none(self):
        """Empty string returns None."""
        assert filter_mod.parse_date_posted("") is None

    def test_parse_whitespace_only_returns_none(self):
        """Whitespace-only string returns None."""
        assert filter_mod.parse_date_posted("   ") is None

    def test_parse_string_none_returns_none(self):
        """String 'none' (case-insensitive) returns None."""
        assert filter_mod.parse_date_posted("none") is None
        assert filter_mod.parse_date_posted("NONE") is None

    def test_parse_string_nan_returns_none(self):
        """String 'nan' returns None."""
        assert filter_mod.parse_date_posted("nan") is None
        assert filter_mod.parse_date_posted("NaN") is None

    def test_parse_string_nat_returns_none(self):
        """String 'NaT' returns None."""
        assert filter_mod.parse_date_posted("NaT") is None
        assert filter_mod.parse_date_posted("nat") is None

    def test_parse_unparseable_returns_none(self):
        """Unparseable string returns None."""
        assert filter_mod.parse_date_posted("yesterday") is None
        assert filter_mod.parse_date_posted("not a date") is None


class TestExtractResumeTextDispatch:
    """filter.extract_resume_text now delegates to pipeline.resume_text, so a
    DOCX/ODT resume scores the same as a PDF."""

    def test_extracts_docx(self, tmp_path):
        from docx import Document
        d = Document()
        d.add_paragraph("Jane Dev")
        d.add_paragraph("Python, AWS, Kubernetes")
        f = tmp_path / "resume.docx"
        d.save(str(f))
        text = filter_mod.extract_resume_text(f)
        assert "Python, AWS, Kubernetes" in text


class TestFindSkillsSection:
    """Test filter.find_skills_section function."""

    def test_find_skills_returns_section_to_next_header(self):
        """Extract text from Skills header to next section header."""
        text = "\nPROFESSIONAL SUMMARY\nSome summary text.\n\nSKILLS\nPython\nJava\nSQL\n\nEXPERIENCE\nSome experience.\n"
        result = filter_mod.find_skills_section(text)
        # Result should include SKILLS header
        assert "SKILLS" in result
        # Should not include EXPERIENCE section
        assert "EXPERIENCE" not in result

    def test_find_skills_no_header_returns_empty(self):
        """No recognized header returns empty string."""
        text = "Just some plain text with no skills section."
        result = filter_mod.find_skills_section(text)
        assert result == ""

    def test_find_skills_at_end_of_document(self):
        """Skills section at end returns from header to doc end."""
        text = """
PROFESSIONAL SUMMARY
Some text.

SKILLS
Python
Java
"""
        result = filter_mod.find_skills_section(text)
        # Just check that we got some content from skills section
        assert len(result) > 5

    def test_find_skills_case_insensitive(self):
        """Recognizes various case-insensitive headers."""
        text1 = "SKILLS\nPython\n\nEXPERIENCE\nStuff"
        result1 = filter_mod.find_skills_section(text1)
        # Should match SKILLS header and include content before EXPERIENCE
        assert len(result1) > 0

        text2 = "Technical Skills:\nJava\n\nEducation\nStuff"
        result2 = filter_mod.find_skills_section(text2)
        # Should match Technical Skills header
        assert len(result2) > 0

    def test_find_skills_competencies_header(self):
        """Recognizes 'Core Competencies' header."""
        text = "Core Competencies:\nDocker\nTerraform\n\nExperience\nStuff"
        result = filter_mod.find_skills_section(text)
        # Should match competencies header and return non-empty result
        assert len(result) > 0

    def test_find_skills_qualified_header(self):
        """'CORE SKILLS' — a qualifier in front of the noun. The old flat
        alternation listed 'core competencies' but not 'core skills', and
        anchored the noun to the start of the line, so this missed silently:
        the section came back empty and the weight-2 boost never reached the
        resume's most discriminating terms (EMR names, tools, certifications)."""
        text = ("CORE SKILLS\n"
                "Systems & Compliance:  Cerner (Oracle Health)  \u2022  eClinicalWorks\n"
                "\nPROFESSIONAL EXPERIENCE\nStuff")
        result = filter_mod.find_skills_section(text)
        assert "Cerner" in result
        assert "PROFESSIONAL EXPERIENCE" not in result

    def test_find_skills_conjoined_header(self):
        """'CERTIFICATIONS & TRAINING' — a conjoined tail. '&' appeared only in
        the next-header lookahead, never in the header pattern itself."""
        text = ("CERTIFICATIONS & TRAINING\n"
                "Basic Life Support (BLS) \u2014 Certified  \u2022  HIPAA Privacy\n"
                "\nEDUCATION\nStuff")
        result = filter_mod.find_skills_section(text)
        assert "BLS" in result
        assert "EDUCATION" not in result

    def test_find_skills_header_variants(self):
        """The qualifier/tail shapes resumes actually use."""
        for header in ("KEY SKILLS", "AREAS OF EXPERTISE", "SKILLS AND ABILITIES",
                       "Technical Skills:", "Proficiencies", "Additional Skills"):
            text = f"{header}\nTools:  Docker  \u2022  Terraform\n\nEXPERIENCE\nStuff"
            result = filter_mod.find_skills_section(text)
            assert "Docker" in result, f"missed {header!r}"

    def test_skills_header_does_not_match_prose_or_compounds(self):
        """The pattern matches a whole line, so it must not fire on a sentence
        that merely starts with the word, nor on a compound like SKILLSET."""
        for line in ("Skills learned on the job", "SKILLSET",
                     "PROFESSIONAL EXPERIENCE", "EDUCATION"):
            text = f"SUMMARY\nx\n\n{line}\nTools:  Docker\n"
            assert filter_mod.find_skills_section(text) == "", f"matched {line!r}"

    @pytest.mark.xfail(
        reason="Pre-existing: the next-header probe r'\\n[A-Z][A-Za-z &/]{3,40}\\n' "
               "matches any bare capitalized line, so a skills section whose "
               "first entry is a single capitalized word ('Cerner', 'Python') "
               "truncates to just the header and yields no tokens. Real resumes "
               "usually escape this by using ':'/'\u2022' delimiters, which the "
               "probe's character class excludes. Distinct from the header-"
               "matching fix; changing the probe risks over-extending a section "
               "into Experience, so it needs its own decision.",
        strict=True,
    )
    def test_bare_capitalized_entry_does_not_truncate_section(self):
        text = "CORE SKILLS\nCerner\neClinicalWorks\n\nPROFESSIONAL EXPERIENCE\nStuff"
        assert "Cerner" in filter_mod.find_skills_section(text)


class TestExtractSkillsSectionTokens:
    """Test filter.extract_skills_section_tokens function."""

    def test_extract_comma_separated(self):
        """Extract comma-separated skills."""
        chunk = "Skills:\nPython, Java, SQL"
        result = filter_mod.extract_skills_section_tokens(chunk)
        assert "python" in result
        assert "java" in result
        assert "sql" in result

    def test_extract_bullet_separated(self):
        """Extract bullet-separated skills."""
        chunk = "Skills:\n• Docker\n• Terraform\n- AWS"
        result = filter_mod.extract_skills_section_tokens(chunk)
        assert "docker" in result
        assert "terraform" in result
        assert "aws" in result

    def test_extract_strips_category_labels(self):
        """Strip category labels like 'Languages:'."""
        chunk = "Skills:\nLanguages: Python, Java\nFrameworks: React, Angular"
        result = filter_mod.extract_skills_section_tokens(chunk)
        assert "python" in result
        assert "java" in result
        assert "react" in result
        assert "languages" not in result
        assert "frameworks" not in result

    def test_extract_filters_noise_words(self):
        """Filter out noise words like 'experience', 'team'."""
        chunk = "Skills:\nexperience, team, Python, work, Java"
        result = filter_mod.extract_skills_section_tokens(chunk)
        assert "python" in result
        assert "java" in result
        assert "experience" not in result
        assert "team" not in result
        assert "work" not in result

    def test_extract_empty_input_returns_empty_set(self):
        """Empty input returns empty set."""
        result = filter_mod.extract_skills_section_tokens("")
        assert result == set()

    def test_extract_too_short_tokens_excluded(self):
        """Tokens shorter than 2 chars are excluded; 2 chars is minimum."""
        chunk = "Skills:\na, Python, JS, Java"
        result = filter_mod.extract_skills_section_tokens(chunk)
        assert "python" in result
        assert "java" in result
        assert "a" not in result
        # JS is exactly 2 chars, which meets the minimum length requirement
        assert "js" in result

    def test_extract_max_length_30_chars(self):
        """Tokens longer than 30 chars are excluded."""
        chunk = "Skills:\nshort_skill, " + "x" * 31 + ", python"
        result = filter_mod.extract_skills_section_tokens(chunk)
        assert "short_skill" in result
        assert "python" in result
        assert "x" * 31 not in result


class TestIsAllNoise:
    """Test filter._is_all_noise function."""

    def test_is_all_noise_all_noise_words(self):
        """All noise words return True."""
        assert filter_mod._is_all_noise("experience team") is True
        assert filter_mod._is_all_noise("work help") is True

    def test_is_all_noise_one_real_word(self):
        """One non-noise word returns False."""
        assert filter_mod._is_all_noise("building rest apis") is False
        assert filter_mod._is_all_noise("experience python") is False

    def test_is_all_noise_empty_string(self):
        """Empty string returns False (no parts)."""
        assert filter_mod._is_all_noise("") is False

    def test_is_all_noise_single_noise_word(self):
        """Single noise word returns True."""
        assert filter_mod._is_all_noise("work") is True
        assert filter_mod._is_all_noise("team") is True


class TestScoreJob:
    """Test filter.score_job function."""

    def test_score_negative_title_returns_none(self):
        """Negative title match returns None."""
        row = {"title": "senior engineer", "description": "", "skills": ""}
        keywords = {}
        result = filter_mod.score_job(row, keywords, [], ["senior"])
        assert result is None

    def test_score_negative_title_word_boundary(self):
        """Word boundary: 'seniority' does NOT match 'senior'."""
        row = {"title": "seniority engineer", "description": "", "skills": ""}
        keywords = {}
        result = filter_mod.score_job(row, keywords, [], ["senior"])
        assert result is not None  # Not excluded
        score, matches = result
        assert score == 0

    def test_score_keyword_match_in_description(self):
        """Keyword in description contributes to score."""
        row = {"title": "developer", "description": "rest apis", "skills": ""}
        keywords = {"rest apis": 2}
        score, matches = filter_mod.score_job(row, keywords, [], [])
        assert score == 2
        assert "rest apis" in matches

    def test_score_keyword_match_skills_boost(self):
        """Keyword in skills field with weight 2."""
        row = {"title": "developer", "description": "", "skills": "spring boot java"}
        keywords = {"spring boot": 2}
        score, matches = filter_mod.score_job(row, keywords, [], [])
        assert score == 2
        assert "spring boot" in matches

    def test_score_target_title_match(self):
        """Target title match adds SCORE_TITLE_MATCH bonus."""
        row = {"title": "software engineer", "description": "", "skills": ""}
        keywords = {}
        score, matches = filter_mod.score_job(row, keywords, ["software engineer"], [])
        assert score == 5  # SCORE_TITLE_MATCH = 5
        assert "title:software engineer" in matches

    def test_score_no_match_returns_zero_score(self):
        """No matches return (0, [])."""
        row = {"title": "developer", "description": "no keywords here", "skills": ""}
        keywords = {"unrelated": 1}
        score, matches = filter_mod.score_job(row, keywords, [], [])
        assert score == 0
        assert matches == []

    def test_score_multiple_keywords_cumulative(self):
        """Multiple keywords add cumulatively."""
        row = {
            "title": "developer",
            "description": "python java",
            "skills": "",
        }
        keywords = {"python": 1, "java": 1}
        score, matches = filter_mod.score_job(row, keywords, [], [])
        assert score == 2
        assert "python" in matches
        assert "java" in matches

    def test_score_negative_title_case_insensitive(self):
        """Negative title matching is case-insensitive."""
        row = {"title": "SENIOR ENGINEER", "description": "", "skills": ""}
        keywords = {}
        result = filter_mod.score_job(row, keywords, [], ["senior"])
        assert result is None

    def test_score_empty_row_no_crash(self):
        """Empty row fields don't crash."""
        row = {"title": "", "description": "", "skills": ""}
        keywords = {}
        score, matches = filter_mod.score_job(row, keywords, [], [])
        assert score == 0
        assert matches == []

    def test_score_none_target_title_no_crash(self):
        """A null/empty entry in target_titles (a bare `-` in YAML parses to
        None) must not crash scoring. `_compile_alternation` already drops
        falsy entries; the target_lookup build must be equally None-tolerant."""
        row = {"title": "research assistant", "description": "", "skills": ""}
        keywords = {}
        score, matches = filter_mod.score_job(
            row, keywords, ["research", None, "", "assistant"], []
        )
        assert score == 10  # two target-title hits × SCORE_TITLE_MATCH (5)
        assert "title:research" in matches
        assert "title:assistant" in matches

    def test_score_none_negative_title_no_crash(self):
        """A null/empty entry in negative_titles must not crash either."""
        row = {"title": "senior engineer", "description": "", "skills": ""}
        result = filter_mod.score_job(row, {}, [], ["senior", None, ""])
        assert result is None


class TestRun:
    """Test filter.run function with mocked pdfplumber."""

    def test_run_filters_and_writes(
        self, jobs_csv, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """run() filters jobs and writes output CSV."""
        jobs_path, output_path = patch_filter_paths

        # Copy jobs_csv to the patched path
        jobs_csv_content = jobs_csv.read_text()
        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text(jobs_csv_content)

        # Set RESUME_PATH env var
        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles:
    - "software engineer"
    - "backend engineer"
  negative_titles:
    - "senior"
  min_score: 1
""")

        filter_mod.run(config)

        # Verify output file was created
        assert output_path.exists()
        content = output_path.read_text()
        # File may be empty if no jobs pass threshold, or may have rows
        if content.strip():
            df = pd.read_csv(output_path)
            assert len(df) >= 1

    def test_run_missing_jobs_csv_raises(self, patch_filter_paths, fake_pdf, monkeypatch):
        """Missing jobs.csv raises FileNotFoundError."""
        jobs_path, output_path = patch_filter_paths
        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: []
  negative_titles: []
  min_score: 5
""")

        with pytest.raises(FileNotFoundError):
            filter_mod.run(config)

    def test_run_missing_resume_raises(self, jobs_csv, patch_filter_paths, monkeypatch):
        """Missing resume PDF raises FileNotFoundError."""
        jobs_path, output_path = patch_filter_paths
        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text(jobs_csv.read_text())

        # Point to nonexistent resume
        monkeypatch.setenv("RESUME_PATH", "/nonexistent/resume.pdf")

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: []
  negative_titles: []
  min_score: 5
""")

        with pytest.raises(FileNotFoundError):
            filter_mod.run(config)

    def test_run_writes_empty_file_when_nothing_passes(
        self, jobs_csv, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """When no jobs pass threshold, writes empty file."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text(jobs_csv.read_text())

        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: ["nonexistent role"]
  negative_titles: []
  min_score: 100
""")

        filter_mod.run(config)

        assert output_path.exists()
        content = output_path.read_text()
        assert content == ""

    def test_run_handles_zero_byte_jobs_csv(
        self, patch_filter_paths, filtered_csv, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """A zero-row scrape truncates jobs.csv to zero bytes (no header line).
        csv.DictReader reads that as zero rows, so the stage no-ops instead of
        raising on the missing header.

        filtered_jobs.csv is seeded first: without a stale file to overwrite,
        "correctly truncated" and "never written" look identical, and the
        staleness this guards against would just move one stage downstream."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text("", encoding="utf-8")
        output_path.write_text(filtered_csv.read_text(encoding="utf-8"), encoding="utf-8")

        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: []
  negative_titles: []
  min_score: 5
""")

        filter_mod.run(config)

        assert output_path.read_text() == ""

    def test_run_handles_header_only_jobs_csv(
        self, patch_filter_paths, filtered_csv, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """The shape the whole contract is about, pinned for filter too.

        bridge and screen both have a header-only case; filter's read semantics
        changed most and had none, so a regression in the `not rows` guard — or
        a return to reading before the emptiness test — would go green here.
        """
        jobs_path, output_path = patch_filter_paths
        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text("title,company,job_url,date_posted\n", encoding="utf-8")
        output_path.write_text(filtered_csv.read_text(encoding="utf-8"), encoding="utf-8")

        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))
        config = jobs_path.parent.parent / "config.yml"
        config.write_text(
            "filter:\n  target_titles: []\n  negative_titles: []\n  min_score: 5\n"
        )

        filter_mod.run(config)

        assert output_path.read_text() == ""

    def test_zero_byte_jobs_csv_does_not_need_a_resume(
        self, patch_filter_paths, filtered_csv, monkeypatch
    ):
        """A missing resume is not a reason to fail a run with nothing to score.

        The resume gate used to sit above the CSV read, so a fork whose
        RESUME_TXT_B64 failed to decode reddened the daily on exactly the days
        the scrape came back empty — two unrelated problems reported as one."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text("", encoding="utf-8")
        output_path.write_text(filtered_csv.read_text(encoding="utf-8"), encoding="utf-8")

        monkeypatch.setenv("RESUME_PATH", "/nonexistent/resume.pdf")

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: []
  negative_titles: []
  min_score: 5
""")

        assert filter_mod.run(config) == output_path
        assert output_path.read_text() == ""

    def test_run_skips_the_resume_work_when_every_row_is_pre_filtered(
        self, jobs_csv, patch_filter_paths, filtered_csv, monkeypatch
    ):
        """The date and location cuts need no keywords, so when they leave
        nothing behind the resume work is skipped too — the truncated-scrape
        case above is just the extreme of it.

        Probed the same way as its sibling, with a resume that isn't there:
        reaching the resume block at all raises, so completing proves it was
        never entered. Every fixture row is dated 2026-05, so a one-hour
        max_age_hours ages all five out."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text(jobs_csv.read_text())
        output_path.write_text(filtered_csv.read_text(encoding="utf-8"), encoding="utf-8")

        monkeypatch.setenv("RESUME_PATH", "/nonexistent/resume.pdf")

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: []
  negative_titles: []
  min_score: 5
  max_age_hours: 1
""")

        assert filter_mod.run(config) == output_path
        assert output_path.read_text() == ""

    def test_run_returns_output_path(
        self, jobs_csv, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """run() returns the output path."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text(jobs_csv.read_text())

        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: []
  negative_titles: []
  min_score: 1
""")

        result = filter_mod.run(config)
        assert result == output_path

    def test_run_max_age_hours_filters_old_jobs(
        self, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """max_age_hours filters out jobs older than N hours."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        old_date = (datetime.now() - timedelta(days=30)).strftime("%Y-%m-%d")
        new_date = datetime.now().strftime("%Y-%m-%d")

        csv_content = (
            "id,job_url,title,company,location,date_posted,description,skills,is_remote\n"
            f'1,https://old.com,engineer,old_co,NYC,{old_date},stuff,,""\n'
            f'2,https://new.com,engineer,new_co,NYC,{new_date},stuff,,""\n'
        )
        jobs_path.write_text(csv_content)

        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  max_age_hours: 168
  target_titles: ["engineer"]
  negative_titles: []
  min_score: 1
""")

        filter_mod.run(config)

        # Only the recent job should be in output
        df = pd.read_csv(output_path)
        assert len(df) == 1
        assert "new.com" in df.iloc[0]["job_url"]

    def test_run_missing_date_kept_regardless(
        self, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """Jobs with blank date_posted are kept even with max_age_hours set."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        recent_date = datetime.now().strftime("%Y-%m-%d")
        csv_content = (
            "id,job_url,title,company,location,date_posted,description,skills,is_remote\n"
            '1,https://old.com,engineer,old_co,NYC,,stuff,,""\n'
            f'2,https://new.com,engineer,new_co,NYC,{recent_date},stuff,,""\n'
        )
        jobs_path.write_text(csv_content)

        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  max_age_hours: 48
  target_titles: ["engineer"]
  negative_titles: []
  min_score: 1
""")

        filter_mod.run(config)

        # Both jobs should be in output (missing date is kept)
        df = pd.read_csv(output_path)
        assert len(df) == 2

    def test_run_keyword_overrides_applied(
        self, jobs_csv, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """keyword_overrides are applied correctly."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        jobs_path.write_text(jobs_csv.read_text())

        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: []
  negative_titles: []
  min_score: 1
  keyword_overrides:
    rest apis: 10
""")

        filter_mod.run(config)

        content = output_path.read_text()
        if content.strip():
            df = pd.read_csv(output_path)
            if len(df) > 0:
                # At least one row mentioning "rest apis" should have high score
                rest_api_rows = df[df["matched_keywords"].str.contains("rest apis", na=False)]
                if len(rest_api_rows) > 0:
                    assert rest_api_rows.iloc[0]["relevance_score"] >= 10

    def test_run_uses_txt_sibling_when_present(
        self, patch_filter_paths, monkeypatch, mocker
    ):
        """When a .txt sibling exists beside the PDF path, it's used without calling pdfplumber."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        csv_content = (
            "id,job_url,title,company,location,date_posted,description,skills,is_remote\n"
            '1,https://job1.com,software engineer,acme,NYC,2026-05-12,python rest apis,,""\n'
        )
        jobs_path.write_text(csv_content)

        # Create a .txt file that lives beside (but not replacing) the PDF path
        fake_pdf = jobs_path.parent / "resume.pdf"
        fake_pdf.write_bytes(b"%PDF-1.4")
        resume_txt = jobs_path.parent / "resume.txt"
        resume_txt.write_text(SYNTHETIC_RESUME, encoding="utf-8")
        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        # pdfplumber should NOT be called — if it is, raise to catch the bug
        mocker.patch("pdfplumber.open", side_effect=AssertionError("pdfplumber should not be called"))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: ["software engineer"]
  negative_titles: []
  min_score: 1
""")

        filter_mod.run(config)
        assert output_path.exists()
        import pandas as pd
        df = pd.read_csv(output_path)
        assert len(df) >= 1

    def test_run_sorts_output_by_score_descending(
        self, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """Output is sorted by relevance_score descending."""
        jobs_path, output_path = patch_filter_paths

        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        csv_content = (
            "id,job_url,title,company,location,date_posted,description,skills,is_remote\n"
            '1,https://job1.com,engineer,a,NYC,2026-05-12,rest apis python,,""\n'
            '2,https://job2.com,engineer,b,NYC,2026-05-12,python,,""\n'
            '3,https://job3.com,software engineer,c,NYC,2026-05-12,stuff,,""\n'
        )
        jobs_path.write_text(csv_content)

        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: ["software engineer"]
  negative_titles: []
  min_score: 1
""")

        filter_mod.run(config)

        df = pd.read_csv(output_path)
        if len(df) > 1:
            # Verify scores are in descending order
            scores = df["relevance_score"].tolist()
            assert scores == sorted(scores, reverse=True)


class TestIsEligible:
    """Test filter.is_eligible — the location/description pre-filter gate."""

    def _pat(self, terms):
        return filter_mod._compile_alternation(terms)

    def test_no_constraints_is_eligible(self):
        """No patterns → always eligible."""
        row = {"location": "Bratislava, Slovakia", "description": "", "is_remote": ""}
        assert filter_mod.is_eligible(row, None, None, None) is True

    def test_short_token_does_not_match_inside_word(self):
        """Word boundaries: negative_locations ['US'] must NOT exclude 'Moscow, Russia'
        just because 'Russia' contains the substring 'us'."""
        row = {"location": "Moscow, Russia", "description": "", "is_remote": ""}
        assert filter_mod.is_eligible(row, self._pat(["US"]), None, None) is True

    def test_short_token_matches_as_whole_word(self):
        """The same ['US'] term still excludes a location where 'US' stands alone."""
        row = {"location": "Dallas, TX, US", "description": "", "is_remote": ""}
        assert filter_mod.is_eligible(row, self._pat(["US"]), None, None) is False

    def test_negative_location_excludes_non_remote(self):
        """A non-remote job in a negative location is excluded."""
        row = {"location": "Bratislava, Slovakia", "description": "", "is_remote": ""}
        assert filter_mod.is_eligible(row, self._pat(["Slovakia"]), None, None) is False

    def test_remote_bypasses_negative_location(self):
        """is_remote=true makes location irrelevant."""
        row = {"location": "Bratislava, Slovakia", "description": "", "is_remote": "true"}
        assert filter_mod.is_eligible(row, self._pat(["Slovakia"]), None, None) is True

    def test_eligible_locations_allowlist_keeps_match(self):
        """A non-remote location containing an allowlisted term is kept."""
        row = {"location": "Dallas, TX, United States", "description": "", "is_remote": ""}
        assert filter_mod.is_eligible(row, None, self._pat(["United States", "USA"]), None) is True

    def test_eligible_locations_allowlist_excludes_non_match(self):
        """A non-remote location matching no allowlist term is excluded."""
        row = {"location": "Berlin, Germany", "description": "", "is_remote": ""}
        assert filter_mod.is_eligible(row, None, self._pat(["United States", "USA"]), None) is False

    def test_empty_location_not_excluded_by_allowlist(self):
        """An unknown/blank location is ambiguous — don't exclude on the allowlist."""
        row = {"location": "", "description": "", "is_remote": ""}
        assert filter_mod.is_eligible(row, None, self._pat(["United States"]), None) is True

    def test_negative_description_term_excludes(self):
        """A description matching a negative term (e.g. a required clearance) is excluded."""
        row = {
            "location": "Dallas, TX",
            "description": "Must hold an active TS/SCI security clearance.",
            "is_remote": "",
        }
        assert filter_mod.is_eligible(row, None, None, self._pat(["security clearance", "ts/sci"])) is False

    def test_no_description_pattern_keeps_job(self):
        """With no negative-description terms configured, the same job is kept (opt-in)."""
        row = {
            "location": "Dallas, TX",
            "description": "Must hold an active TS/SCI security clearance.",
            "is_remote": "",
        }
        assert filter_mod.is_eligible(row, None, None, None) is True

    def test_description_term_excludes_even_when_remote(self):
        """Description terms are location-independent — a remote match is still excluded."""
        row = {
            "location": "Anywhere",
            "description": "Requires an active Secret clearance.",
            "is_remote": "true",
        }
        assert filter_mod.is_eligible(row, None, None, self._pat(["secret clearance"])) is False

    def test_description_term_matches_exact_substring_only(self):
        """A loose mention ('clearance of priorities') doesn't match the configured terms."""
        row = {
            "location": "Dallas, TX",
            "description": "We value clear communication and a clearance of priorities.",
            "is_remote": "",
        }
        assert filter_mod.is_eligible(row, None, None, self._pat(["security clearance", "ts/sci"])) is True


class TestRunEligibility:
    """Integration: location pre-filter wired through filter.run()."""

    def test_run_excludes_negative_location_keeps_remote(
        self, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """Non-remote overseas job dropped; remote job in the same place survives."""
        jobs_path, output_path = patch_filter_paths
        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        csv_content = (
            "id,job_url,title,company,location,date_posted,description,skills,is_remote\n"
            '1,https://intl.com,engineer,a,"Bratislava, Slovakia",2026-05-12,stuff,,""\n'
            '2,https://remote.com,engineer,b,"Bratislava, Slovakia",2026-05-12,stuff,,"true"\n'
            '3,https://us.com,engineer,c,"Dallas, TX",2026-05-12,stuff,,""\n'
        )
        jobs_path.write_text(csv_content)
        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: ["engineer"]
  negative_titles: []
  min_score: 1
  negative_locations: ["Slovakia"]
""")
        filter_mod.run(config)

        urls = set(pd.read_csv(output_path)["job_url"])
        assert "https://intl.com" not in urls
        assert "https://remote.com" in urls
        assert "https://us.com" in urls

    def test_run_eligible_locations_allowlist(
        self, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """Only non-remote jobs in an allowlisted location are kept."""
        jobs_path, output_path = patch_filter_paths
        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        csv_content = (
            "id,job_url,title,company,location,date_posted,description,skills,is_remote\n"
            '1,https://us.com,engineer,a,"Austin, TX, United States",2026-05-12,stuff,,""\n'
            '2,https://de.com,engineer,b,"Berlin, Germany",2026-05-12,stuff,,""\n'
        )
        jobs_path.write_text(csv_content)
        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: ["engineer"]
  negative_titles: []
  min_score: 1
  eligible_locations: ["United States", "USA"]
""")
        filter_mod.run(config)

        urls = set(pd.read_csv(output_path)["job_url"])
        assert urls == {"https://us.com"}

    def test_run_negative_description_terms(
        self, patch_filter_paths, fake_pdf, monkeypatch, mock_pdf_extract
    ):
        """Jobs whose description contains a configured negative term are dropped."""
        jobs_path, output_path = patch_filter_paths
        jobs_path.parent.mkdir(parents=True, exist_ok=True)
        csv_content = (
            "id,job_url,title,company,location,date_posted,description,skills,is_remote\n"
            '1,https://clear.com,engineer,a,"Dallas, TX",2026-05-12,"Requires TS/SCI security clearance",,""\n'
            '2,https://ok.com,engineer,b,"Dallas, TX",2026-05-12,"Build web apps",,""\n'
        )
        jobs_path.write_text(csv_content)
        monkeypatch.setenv("RESUME_PATH", str(fake_pdf))

        config = jobs_path.parent.parent / "config.yml"
        config.write_text("""
filter:
  target_titles: ["engineer"]
  negative_titles: []
  min_score: 1
  negative_description_terms: ["security clearance"]
""")
        filter_mod.run(config)

        urls = set(pd.read_csv(output_path)["job_url"])
        assert urls == {"https://ok.com"}
