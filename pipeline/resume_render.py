"""Deterministic résumé renderer (Commit 3b): a generic content-JSON → a one-page
.docx via python-docx, with tight base typography (US Letter, 0.6"/0.8"/0.5"
margins, 9.5pt Arial body, 10.5pt ruled headings) — the Cowork settings that fill
92–96% of a page for content-rich résumés.

`scale` resizes the WHOLE layout (fonts + spacing + line height) together, so the
same content can be grown to fill a page (lighter résumés) or kept tight (denser
ones). resume_build.fit_to_page searches the scale that fills one page; the section
set is generic and user-agnostic so it renders any candidate's content.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

FONT = "Arial"
# Base point sizes mirror the Cowork resume.js that reliably fills one page; `scale`
# multiplies them and every gap, and line spacing opens up with it too (never below
# single) — so leading grows faster than font size above scale 1.0. That combined
# effect, not font size alone, is what fills the page (see _Layout._para).
NAME_PT, CONTACT_PT, HEAD_PT, ROLE_PT, BODY_PT = 15, 8.5, 10.5, 9, 9.5
PAGE_W, PAGE_H = Twips(12240), Twips(15840)                     # US Letter
M_TOP, M_SIDE, M_BOTTOM = Twips(864), Twips(1152), Twips(720)   # 0.6 / 0.8 / 0.5 in
CONTENT_W = Twips(12240 - 1152 - 1152)                          # right tab stop for dates/loc

DEFAULT_PROJECTS_HEADING = "Projects"
DEFAULT_EDUCATION_HEADING = "Education"


def _s(v) -> str:
    """A field's text, tolerating a missing / null / non-string value — an LLM-fed
    content-JSON can carry `null` for any field, which a `.get(k, "")` default (only
    triggered by an ABSENT key) would otherwise pass straight into a string concat."""
    return "" if v is None else str(v)


def _list(v) -> list:
    """A field's list, tolerating a null / non-list value (e.g. `"bullets": null`)."""
    return v if isinstance(v, list) else []


def _bottom_border(paragraph):
    """A thin rule under a section heading (python-docx exposes no border API)."""
    pPr = paragraph._p.get_or_add_pPr()
    borders, bottom = OxmlElement("w:pBdr"), OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "4"), ("w:space", "1"), ("w:color", "000000")):
        bottom.set(qn(k), v)
    borders.append(bottom)
    pPr.append(borders)


class _Layout:
    """Builds paragraphs at a uniform `scale`: fonts and gaps multiply by scale,
    and line spacing grows with it (never below single, to avoid clipping)."""

    def __init__(self, doc, scale: float):
        self.doc, self.s = doc, scale

    def _font(self, run, size_pt, *, bold=False, italic=False):
        run.font.name = FONT
        run.font.size = Pt(size_pt * self.s)
        run.font.bold, run.font.italic = bold, italic

    def _para(self, *, before=0, after=4, align=None):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before, pf.space_after = Pt(before * self.s), Pt(after * self.s)
        pf.line_spacing = max(1.0, self.s)
        if align is not None:
            p.alignment = align
        return p

    def heading(self, text):
        p = self._para(before=6, after=3)
        self._font(p.add_run(text.upper()), HEAD_PT, bold=True)
        _bottom_border(p)

    def centered(self, text, size_pt, *, bold=False, after):
        p = self._para(after=after, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._font(p.add_run(_s(text)), size_pt, bold=bold)

    def body(self, text, *, after=3):
        self._font(self._para(after=after).add_run(_s(text)), BODY_PT)

    def skills_row(self, row):
        p = self._para(after=1)
        self._font(p.add_run(_s(row.get("label")) + ": "), BODY_PT, bold=True)
        self._font(p.add_run(_s(row.get("items"))), BODY_PT)

    def bullet(self, text):
        p = self._para(after=1)
        p.paragraph_format.left_indent = Twips(180)
        p.paragraph_format.first_line_indent = Twips(-180)   # hanging indent → aligned wrap
        self._font(p.add_run("• " + _s(text)), BODY_PT)

    def entry(self, e, *, with_meta):
        """org (bold) [tab] dates, then role (italic) [tab] loc, then bullets.
        Projects pass with_meta=False (no dates/loc)."""
        p1 = self._para(after=1)
        p1.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
        self._font(p1.add_run(_s(e.get("org"))), BODY_PT, bold=True)
        if with_meta and e.get("dates"):
            self._font(p1.add_run("\t" + _s(e.get("dates"))), BODY_PT)
        if e.get("role") or (with_meta and e.get("loc")):
            p2 = self._para(after=1)
            p2.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
            self._font(p2.add_run(_s(e.get("role"))), ROLE_PT, italic=True)
            if with_meta and e.get("loc"):
                self._font(p2.add_run("\t" + _s(e.get("loc"))), ROLE_PT, italic=True)
        for b in _list(e.get("bullets")):
            if b:
                self.bullet(b)


def render_docx(content: dict, out_path, scale: float = 1.0) -> Path:
    """Render a content-JSON to a one-page-styled .docx at out_path. `scale`
    resizes the whole layout. Every section is optional — an empty one is omitted
    (no dangling heading). The `projects` flex section renders after experience, or
    before it when `projects_first`."""
    out_path = Path(out_path)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = PAGE_W, PAGE_H
    sec.top_margin, sec.bottom_margin = M_TOP, M_BOTTOM
    sec.left_margin, sec.right_margin = M_SIDE, M_SIDE
    L = _Layout(doc, scale)

    L.centered(content.get("name", ""), NAME_PT, bold=True, after=1)
    if content.get("contact"):
        L.centered(content["contact"], CONTACT_PT, after=4)

    if content.get("summary"):
        L.heading("Professional Summary")
        L.body(content["summary"])

    if content.get("skills"):
        L.heading("Skills")
        for row in _list(content.get("skills")):
            L.skills_row(row)

    # Experience then the projects flex section — order swapped by projects_first.
    sections = [("experience", "Professional Experience", True),
                ("projects", content.get("projects_heading") or DEFAULT_PROJECTS_HEADING, False)]
    if content.get("projects_first"):
        sections.reverse()
    for key, heading, with_meta in sections:
        if content.get(key):
            L.heading(heading)
            for e in _list(content.get(key)):
                L.entry(e, with_meta=with_meta)

    if content.get("education"):
        L.heading(content.get("education_heading") or DEFAULT_EDUCATION_HEADING)
        for item in _list(content.get("education")):
            if item:
                L.bullet(item)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
