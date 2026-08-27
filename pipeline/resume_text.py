"""Resume text extraction — format-dispatching over PDF / DOCX / ODT / TXT.

A single entry point both the keyword-scoring filter (pipeline/filter.py) and the
UI onboarding upload (pipeline/app/onboard.py) go through, so every resume format
is read identically. Extracted text is what feeds YAKE keyword scoring and the
cloud evaluator (resumes/resume.txt); the source format doesn't otherwise matter.

DOCX/ODT are preferred over PDF for import: resume tailoring
(pipeline/resume_tailor.py) slot-edits a DOCX directly, and editable formats give
cleaner text than a PDF's reconstructed layout. All four formats work for scoring.

Per-format libraries are imported lazily inside each extractor so importing this
module stays cheap (the UI's onboard module imports it at the top) and a missing
optional dep only bites the user who actually hands us that format.
"""

import sys
import tempfile
from pathlib import Path

from pipeline.stdio import line_buffer_stdout

# What the setup/UI offer as importable resume formats (no .txt — that's the
# generated sidecar, not something a user uploads).
IMPORT_SUFFIXES = (".pdf", ".docx", ".odt")


# pdfplumber inserts a space between two glyphs when the gap between them
# exceeds x_tolerance, which defaults to 3 (points). A PDF whose kerning is
# tighter than that emits no spaces at all — "Managedhigh-volumeinboundcalls" —
# and nothing raises. YAKE then extracts keywords like "microsoftoffice" that
# match no job description ever, every job scores ~0, and the run presents as
# "nothing matched" — which is what a bad search config and a rate-limited
# scrape also look like. So detect that signature and retry tighter.
_PDF_RETRY_X_TOLERANCE = 1.5

# A run-together line yields tokens many times longer than any English word.
# The threshold sits well above the longest tokens a real resume produces
# (URLs, emails, "telecommunications") so a normal PDF never triggers a retry.
_RUNON_TOKEN_LEN = 25
_RUNON_TOKEN_SHARE = 0.05


def _looks_run_together(text: str) -> bool:
    """True when extraction produced implausibly long whitespace-separated
    tokens — the signature of glyph gaps falling under x_tolerance."""
    tokens = text.split()
    if len(tokens) < 20:
        # Too little text to judge. A near-empty extraction is its own problem
        # and callers already surface it; guessing at spacing from a handful of
        # tokens would only add a second failure mode.
        return False
    runon = sum(1 for t in tokens if len(t) > _RUNON_TOKEN_LEN)
    return runon / len(tokens) > _RUNON_TOKEN_SHARE


def _from_pdf(path: Path) -> str:
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        pages = list(pdf.pages)
        text = "\n".join((p.extract_text() or "") for p in pages)
        if not _looks_run_together(text):
            return text
        retry = "\n".join(
            (p.extract_text(x_tolerance=_PDF_RETRY_X_TOLERANCE) or "")
            for p in pages
        )
        # Keep the retry only if it actually resolved the run-together text. A
        # tighter tolerance can split words that were fine, so a retry that
        # still looks wrong is no evidence it helped — prefer the default.
        if _looks_run_together(retry):
            return text
        print(
            f"[resume] {path.name}: words ran together at pdfplumber's default "
            f"spacing; re-extracted at x_tolerance={_PDF_RETRY_X_TOLERANCE}."
        )
        return retry


def _from_docx(path: Path) -> str:
    # Paragraphs plus table cells — skills/tech are frequently laid out in a
    # borderless table, which doc.paragraphs alone would miss.
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _from_odt(path: Path) -> str:
    # Walk the body in document order (rather than getElementsByType, which
    # groups by type and would scramble the name-then-sections ordering the
    # contact-info parser relies on). Headings (text:h) and paragraphs (text:p)
    # are emitted as lines; teletype.extractText flattens nested spans/links.
    from odf.opendocument import load
    from odf import teletype
    doc = load(str(path))
    lines: list[str] = []

    def walk(node) -> None:
        for child in getattr(node, "childNodes", []):
            qname = getattr(child, "qname", None)
            if qname and qname[1] in ("p", "h"):
                lines.append(teletype.extractText(child))
            else:
                walk(child)

    walk(doc.text)
    return "\n".join(lines)


def _from_txt(path: Path) -> str:
    return Path(path).read_text(encoding="utf-8")


def _extractor_for(suffix: str):
    """The extractor for a file suffix, or a clear ValueError so callers can
    surface "use PDF/DOCX/ODT" rather than a cryptic library error. Dispatch is
    by name through module globals, so tests can monkeypatch an extractor.
    `.txt` isn't an import format users pick, but the filter reads a resume.txt
    sidecar directly, so the dispatcher handles it too."""
    name = {".pdf": "_from_pdf", ".docx": "_from_docx",
            ".odt": "_from_odt", ".txt": "_from_txt"}.get(suffix)
    if name is None:
        raise ValueError(
            f"Unsupported resume format {suffix or '(none)'!r}. "
            f"Use one of: {', '.join(IMPORT_SUFFIXES)}."
        )
    return globals()[name]


# Formats extract_resume_text accepts (import formats + the .txt sidecar).
SUPPORTED_SUFFIXES = IMPORT_SUFFIXES + (".txt",)


def extract_resume_text(path: Path) -> str:
    """Extract plain text from a resume file, dispatching on its suffix."""
    path = Path(path)
    return _extractor_for(path.suffix.lower())(path)


def extract_resume_bytes(data: bytes, filename: str) -> str:
    """Extract text from in-memory resume bytes, dispatching on `filename`'s
    suffix (an upload is bytes + a client-supplied name, not a path on disk).

    Writes a temp file with the right suffix so the per-format libraries — which
    open by path — work unchanged."""
    suffix = Path(filename).suffix.lower()
    _extractor_for(suffix)   # reject unsupported formats BEFORE the temp write
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
        f.write(data)
        tmp = Path(f.name)
    try:
        return extract_resume_text(tmp)
    finally:
        try:
            tmp.unlink()
        except OSError:
            pass


def _main(argv: list[str]) -> int:
    """`python -m pipeline.resume_text <path>` — print extracted text to stdout.

    The Node setup script (setup-profile.mjs) shells out to this for DOCX/ODT so
    there's a single extraction implementation and no extra npm dependencies."""
    if len(argv) != 1:
        print("usage: python -m pipeline.resume_text <resume-file>", file=sys.stderr)
        return 2
    try:
        sys.stdout.write(extract_resume_text(Path(argv[0])))
        return 0
    except (ValueError, OSError) as e:
        print(str(e), file=sys.stderr)
        return 1


if __name__ == "__main__":
    line_buffer_stdout()

    raise SystemExit(_main(sys.argv[1:]))
